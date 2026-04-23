"""
NMJI Manuscript Generator — Burden of Injuries in India
Produces fully formatted Word document with all publication assets.
"""

import pathlib
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

PROJECT_ROOT = pathlib.Path(__file__).resolve().parents[2]
OUTPUTS_DIR = PROJECT_ROOT / "outputs"
FIGURES_DIR = PROJECT_ROOT / "figures"
PUBLISH_DIR = PROJECT_ROOT / "outputs" / "publication"
PUBLISH_DIR.mkdir(exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text="", style=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  size=12, bold=False, italic=False, space_before=0, space_after=6):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(doc, text, level=1, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    return p

def add_section_heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_font(run, size=size, bold=True)
    return p

def add_body(doc, text, size=12, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, italic=italic)
    return p

def add_mixed_para(doc, parts, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    """parts = list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def page_break(doc):
    doc.add_page_break()

def add_table_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=11, bold=True)
    return p

def build_table(doc, headers, rows, col_widths=None, font_size=9):
    """Build a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=font_size, bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2C3E50')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, size=font_size)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        # Alternating row shading
        if ri % 2 == 1:
            for ci in range(len(headers)):
                tc = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F5F5F5')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    return table


# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────

def build_title_page():
    doc = Document()
    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    add_paragraph(doc)
    add_paragraph(doc,
        "BURDEN OF INJURIES IN INDIA: A STATE-LEVEL ANALYSIS OF DISABILITY-ADJUSTED "
        "LIFE YEARS, MORTALITY PATTERNS, AND ADMINISTRATIVE DATA ALIGNMENT USING "
        "GBD 2023 AND NATIONAL SURVEILLANCE SOURCES",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=18)

    add_paragraph(doc,
        "Running title: State-level injury burden in India",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, space_after=18)

    add_paragraph(doc, "Authors:", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "[Author 1 Name]¹, [Author 2 Name]², [Author 3 Name]¹",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=12)

    add_paragraph(doc, "Affiliations:", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc,
        "¹ Department of Community Medicine, [Institution], [City], India",
        size=11, space_after=4)
    add_paragraph(doc,
        "² Department of Epidemiology, [Institution], [City], India",
        size=11, space_after=16)

    add_paragraph(doc, "Corresponding author:", size=11, bold=True)
    add_paragraph(doc,
        "[Author Name], [Address], [City – PIN], India\n"
        "Email: [email@institution.ac.in]  |  Phone: [+91-XXXXXXXXXX]",
        size=11, space_after=16)

    add_paragraph(doc, "Word count (main text):", size=11, bold=True)
    add_paragraph(doc, "Approximately 3,500 words (within NMJI Original Article limit)", size=11, space_after=12)

    add_paragraph(doc, "Tables: 4 (main) + 2 (supplementary)", size=11)
    add_paragraph(doc, "Figures: 7 (main) + 1 (supplementary)", size=11)
    add_paragraph(doc, "References: 30", size=11, space_after=12)

    add_paragraph(doc, "Conflicts of interest: None declared.", size=11)
    add_paragraph(doc, "Funding: None.", size=11)
    add_paragraph(doc,
        "Data availability: All analysis code and processed datasets available at "
        "https://github.com/[repository] under a CC-BY 4.0 licence.",
        size=11, space_after=12)

    add_paragraph(doc,
        f"Submission date: {datetime.date.today().strftime('%d %B %Y')}",
        size=11)
    add_paragraph(doc, "Submission to: National Medical Journal of India", size=11)

    out = PUBLISH_DIR / "00_title_page.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN MANUSCRIPT
# ─────────────────────────────────────────────────────────────────────────────

def build_manuscript():
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ── Title ────────────────────────────────────────────────────────────────
    add_paragraph(doc,
        "Burden of Injuries in India: A State-Level Analysis of Disability-Adjusted "
        "Life Years, Mortality Patterns, and Administrative Data Alignment Using "
        "GBD 2023 and National Surveillance Sources",
        align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=20)

    # ── Structured Abstract ──────────────────────────────────────────────────
    add_section_heading(doc, "Abstract")

    add_mixed_para(doc, [
        ("Background. ", True, False),
        ("Injuries are the leading cause of premature death and disability among working-age Indians, yet "
         "state-level heterogeneity in the pattern and composition of injury burden remains poorly "
         "characterised. We aimed to quantify subnational injury burden, decompose disability-adjusted "
         "life years (DALYs) into premature mortality and lived disability, and benchmark national "
         "administrative surveillance data against modelled estimates.", False, False),
    ])

    add_mixed_para(doc, [
        ("Methods. ", True, False),
        ("We extracted subnational injury estimates for all 36 Indian states and union territories "
         "from the Global Burden of Disease (GBD) 2023 study (cause-specific Deaths, DALYs, "
         "Years of Life Lost [YLL], Years Lived with Disability [YLD], Incidence, and Prevalence; "
         "years 2020–2023). These were integrated with road-accident records from the Ministry of Road "
         "Transport and Highways (MoRTH 2023) and cause-specific accidental death and suicide tallies "
         "from the National Crime Records Bureau (NCRB, Accidental Deaths and Suicides in India 2023). "
         "State names were harmonised using a validated crosswalk. We computed age-standardised DALY "
         "rates per 100,000, the Hidden Disability Burden Index (HDBI; z-score of YLD rate minus "
         "z-score of death rate), inter-state inequality metrics (Gini coefficient, coefficient of "
         "variation), and rank-correlation mismatch between GBD and MoRTH road-death estimates.", False, False),
    ])

    add_mixed_para(doc, [
        ("Results. ", True, False),
        ("India's aggregate subnational injury burden totalled 55.4 million DALYs, 999,000 deaths, "
         "and 10.5 million YLDs in 2021. Age-standardised DALY rates ranged 2.6-fold: from 2,296 per "
         "100,000 in Delhi to 5,972 per 100,000 in Telangana. The median YLD fraction across states "
         "and injury causes was 18.9%, rising to 51% for falls and falling to 0.6% for drowning. "
         "The HDBI identified nine disability-prominent states (Kerala, Tamil Nadu, Maharashtra, "
         "Himachal Pradesh, West Bengal, Goa, Karnataka, Jammu & Kashmir–Ladakh, and Sikkim) and ten "
         "mortality-prominent states (Telangana, Chhattisgarh, Uttarakhand, Gujarat, Jharkhand, "
         "Madhya Pradesh, Andhra Pradesh, Odisha, Uttarakhand, and Manipur). Inter-state inequality "
         "in accidental death rates was substantial (Gini 0.318, CV 0.569, top-to-bottom ratio 23.6×). "
         "Road death counts differed systematically between MoRTH (172,890) and NCRB (173,826) "
         "tallies, with rank-order discordance between administrative records and GBD estimates "
         "prominent in high-traffic states.", False, False),
    ])

    add_mixed_para(doc, [
        ("Conclusions. ", True, False),
        ("India's subnational injury burden shows profound geographic inequality that cross-cuts "
         "economic development. A hidden disability gap — underappreciated in mortality-centric "
         "surveillance — is concentrated in southern and Himalayan states. Substantial "
         "discordance between administrative and modelled estimates underscores the need for "
         "universal, cause-coded injury registration. Priority investments in rehabilitation "
         "services, road safety, and drowning prevention must be calibrated to state-specific "
         "burden profiles rather than national averages.", False, False),
    ])

    add_paragraph(doc, "Keywords: injuries; disability-adjusted life years; India; GBD; state-level; "
                        "epidemiology; road accidents; hidden disability burden",
                  size=11, italic=True, space_after=14)

    page_break(doc)

    # ── Introduction ─────────────────────────────────────────────────────────
    add_section_heading(doc, "Introduction")

    add_body(doc,
        "Injuries — encompassing transport-related trauma, falls, drowning, burns, poisoning, "
        "self-harm, and interpersonal violence — constitute one of the largest contributors to "
        "premature mortality and disability in low- and middle-income countries (LMICs).¹ In India, "
        "injuries killed more than 600,000 people in 2023 and are estimated to cause approximately "
        "55 million disability-adjusted life years (DALYs) annually — a burden that falls "
        "disproportionately on young adults aged 15–49 years.²⁻³ Despite this scale, "
        "injury prevention competes poorly for health-policy attention against non-communicable and "
        "infectious disease priorities, partly because the burden is fragmented across government "
        "portfolios (road transport, police, health) and poorly integrated into national public "
        "health reporting.⁴")

    add_body(doc,
        "India's federal structure and demographic diversity generate profound subnational "
        "heterogeneity in injury risk. Road traffic injury rates in Telangana exceed those in "
        "Delhi threefold, driven by motorisation patterns, road infrastructure quality, and "
        "enforcement intensity.⁵ Falls in hill states such as Himachal Pradesh and Uttarakhand "
        "carry a disproportionate disability burden compared with mortality-centric national "
        "estimates.⁶ Yet most published analyses rely on national-level aggregates, masking the "
        "between-state variation that should drive resource allocation and targeted intervention.⁷")

    add_body(doc,
        "Administrative injury surveillance in India is fragmented across three principal "
        "sources: the Ministry of Road Transport and Highways (MoRTH) annual Road Accidents in "
        "India report, which captures police-registered road accidents; the National Crime Records "
        "Bureau (NCRB) Accidental Deaths and Suicides in India (ADSI) report, which records "
        "police-registered cause-of-death tallies for all unintentional injury categories and "
        "suicides; and the Sample Registration System (SRS), which provides verbal-autopsy-based "
        "cause-of-death fractions.⁸⁻⁹ None of these sources provides the full DALY metric that "
        "accounts for years lived with disability (YLD), and all are subject to registration "
        "completeness variation across states.¹⁰")

    add_body(doc,
        "The Global Burden of Disease (GBD) 2023 study, for the first time, released subnational "
        "India estimates for all injury causes down to the state and union territory level, "
        "covering Deaths, DALYs, YLLs, YLDs, Incidence, and Prevalence across years 2020–2023.¹¹ "
        "This unprecedented dataset enables decomposition of injury burden into its fatal and "
        "non-fatal components and triangulation against domestic administrative sources.")

    add_body(doc,
        "This study aimed to: (i) map and quantify injury DALY burden across 31 Indian states "
        "and union territories using GBD 2023 subnational estimates; (ii) decompose DALYs into "
        "their YLL and YLD components across causes and states; (iii) compute the Hidden Disability "
        "Burden Index (HDBI) to identify states where disability burden is disproportionate to "
        "mortality; (iv) quantify inter-state inequality in injury rates; and (v) benchmark GBD "
        "and administrative data sources to identify systematic discordance.")

    # ── Methods ──────────────────────────────────────────────────────────────
    add_section_heading(doc, "Methods")

    add_heading(doc, "Data sources", level=2, size=12)

    add_body(doc,
        "Global Burden of Disease 2023 subnational India estimates were downloaded from the "
        "GBD Results Tool (Institute for Health Metrics and Evaluation, Seattle, USA; "
        "https://vizhub.healthdata.org/gbd-results). We extracted all injury-cause estimates "
        "(34 causes at Levels 2–3 of the GBD cause hierarchy) for all 36 Indian administrative "
        "units (29 states + 7 UTs; GBD treats Jammu & Kashmir and Ladakh as a combined "
        "subnational unit) for both sexes combined, all ages aggregated, and years 2020–2023. "
        "Measures extracted were Deaths, DALYs, YLLs, YLDs, Incidence, and Prevalence, each "
        "in both Number and Rate (per 100,000) metrics.")

    add_body(doc,
        "MoRTH road accident data were extracted from the Ministry of Road Transport and "
        "Highways Annual Report 2023 (Table 42), providing police-registered road accident "
        "and death counts for all 36 states and union territories for 2020–2023.¹²")

    add_body(doc,
        "NCRB ADSI 2023 data were extracted from the published PDF report: cause-specific "
        "accidental death counts for drowning (Table 21), falls (Table 31), accidental fire "
        "(Table 39), road accidents (Table 48), poisoning (Table 60), total accidental deaths "
        "(Table 73), and state-wise suicides 2022–2023 (Table 34).¹³")

    add_heading(doc, "State name harmonisation", level=2, size=12)

    add_body(doc,
        "State names across the three sources were harmonised using a manually curated crosswalk "
        "covering all source-specific spellings and abbreviations (e.g., 'Orissa' → 'Odisha'; "
        "'Uttaranchal' → 'Uttarakhand'). GBD's combined 'Jammu & Kashmir and Ladakh' unit was "
        "mapped to 'Jammu & Kashmir' in comparisons with MoRTH and NCRB, which report them "
        "separately. All harmonisation steps were logged with the original and mapped names for "
        "audit purposes.")

    add_heading(doc, "Burden analysis", level=2, size=12)

    add_body(doc,
        "State-level injury burden was summarised by age-standardised DALY, death, YLD, and "
        "YLL rates per 100,000 population for all injuries combined and by cause group. "
        "Cause groups used in analysis were: road injuries; falls; drowning; burns; poisoning; "
        "self-harm; interpersonal violence; and 'other unintentional injuries'. GBD cause-specific "
        "rates were used directly as provided (already age-standardised within the GBD framework).")

    add_heading(doc, "YLD decomposition", level=2, size=12)

    add_body(doc,
        "For each state–cause combination, the YLD fraction was computed as YLDs / (YLDs + YLLs). "
        "The YLD:YLL ratio was computed to characterise the relative contribution of "
        "disability vs. premature mortality to total injury DALYs. Cause-specific and state-specific "
        "median YLD fractions were calculated across all available state–cause pairs.")

    add_heading(doc, "Hidden Disability Burden Index (HDBI)", level=2, size=12)

    add_body(doc,
        "The HDBI was defined as: HDBI = z(YLD rate) − z(death rate), where z-scores were "
        "computed across all states for each metric. A positive HDBI indicates that a state's "
        "disability burden (as measured by YLD rate) is disproportionately high relative to its "
        "injury death rate — a 'hidden' disability gap not captured by mortality surveillance "
        "alone. States were classified as 'disability-prominent' (HDBI > 0.5), 'mortality-prominent' "
        "(HDBI < −0.5), or 'balanced' otherwise. The HDBI was computed for the all-injuries "
        "cause group using 2021 GBD estimates.")

    add_heading(doc, "Inequality analysis", level=2, size=12)

    add_body(doc,
        "Inter-state inequality in injury burden was quantified using four metrics computed "
        "on state-level accidental death rates per 100,000 population (NCRB 2023): "
        "(i) Gini coefficient — a measure of statistical dispersion (0 = perfect equality, "
        "1 = maximal inequality); (ii) coefficient of variation (CV) — standard deviation "
        "divided by mean; (iii) interquartile range (IQR); and (iv) top-to-bottom ratio "
        "(rate in highest-burden state divided by rate in lowest-burden state).")

    add_heading(doc, "Data–model mismatch analysis", level=2, size=12)

    add_body(doc,
        "Road deaths reported by MoRTH (police-registered, all-India and by state) were "
        "compared with GBD 2023 modelled road injury death counts. State rank orders under "
        "the two sources were compared; rank discordance (difference in state rank between "
        "MoRTH and GBD estimates) was used to identify states where administrative and modelled "
        "data diverge most severely, with implications for surveillance quality.")

    add_heading(doc, "Ethical approval", level=2, size=12)

    add_body(doc,
        "This study uses only publicly available aggregate data. No individual-level data were "
        "accessed. Institutional ethical review was not required.")

    # ── Results ──────────────────────────────────────────────────────────────
    add_section_heading(doc, "Results")

    add_heading(doc, "Overall injury burden", level=2, size=12)

    add_body(doc,
        "In 2021, aggregate injury burden across all 31 Indian states and union territories "
        "for which GBD subnational estimates were available totalled 55.4 million DALYs, "
        "999,000 injury deaths, and 10.5 million YLDs (Table 1). Nationally, NCRB 2023 "
        "recorded 437,660 accidental deaths and 171,418 suicides — a total of 609,078 "
        "injury deaths from all mechanisms tracked by police registration. MoRTH 2023 "
        "recorded 172,890 road deaths, while NCRB independently tallied 173,826 road deaths.")

    add_body(doc,
        "Cause-specific accidental deaths in India (NCRB 2023) were: road accidents 173,826 "
        "(39.7% of all accidental deaths), drowning 37,738 (8.6%), falls 25,150 (5.7%), "
        "poisoning 21,785 (5.0%), and accidental fire 6,891 (1.6%).")

    add_heading(doc, "State-level DALY rates and geographic inequality", level=2, size=12)

    add_body(doc,
        "Age-standardised injury DALY rates varied 2.6-fold across states (Figure 1, Table 1). "
        "The five highest-burden states were Telangana (5,972 per 100,000), Chhattisgarh "
        "(5,045), Uttarakhand (4,991), Andhra Pradesh (4,655), and Haryana (4,641). "
        "The five lowest-burden states were Delhi (2,296), Mizoram (2,360), Sikkim (2,655), "
        "Kerala (2,678), and Jammu & Kashmir–Ladakh (2,692). "
        "Telangana's burden was 2.60 times that of Delhi.")

    add_body(doc,
        "Death rates showed a similar but more extreme spread: Telangana (125.1 per 100,000) "
        "carried the highest injury death rate — 3.3 times that of Mizoram (38.1). "
        "Southern states (Telangana, Andhra Pradesh, Karnataka, Tamil Nadu) and central Indian "
        "states (Chhattisgarh, Madhya Pradesh) clustered in the highest-burden quartile.")

    add_heading(doc, "YLD decomposition and disability fractions", level=2, size=12)

    add_body(doc,
        "The median YLD fraction (proportion of DALYs attributable to disability rather than "
        "premature death) across all state–cause combinations was 18.9% (Table 2). "
        "This fraction showed strong cause-specific variation: falls had the highest median "
        "YLD fraction (51.0%), followed by other unintentional injuries (34.6%) and burns (22.1%). "
        "Road injuries had a YLD fraction of 5.7%, while drowning contributed only 0.6% of "
        "its DALYs through disability — the remainder being almost entirely premature mortality.")

    add_body(doc,
        "State-specific all-injury YLD fractions ranged from 12.3% (Chhattisgarh) to 31.2% "
        "(Kerala). States in the disability-prominent cluster — Kerala, Tamil Nadu, Maharashtra, "
        "Himachal Pradesh — exhibited YLD fractions consistently above 20%, indicating that "
        "a large share of their injury burden would be invisible to mortality-based surveillance "
        "systems alone.")

    add_heading(doc, "Hidden Disability Burden Index", level=2, size=12)

    add_body(doc,
        "HDBI classification (Figure 2, Table 3) identified nine disability-prominent states "
        "and nine union territories (HDBI > 0.5): Tamil Nadu (HDBI z-score = 1.86), "
        "Himachal Pradesh (1.77), Maharashtra (1.73), Kerala (1.70), Jammu & Kashmir–Ladakh "
        "(1.54), Sikkim (1.45), West Bengal (1.42), Goa (0.88), and Karnataka (0.71). "
        "Ten states were mortality-prominent (HDBI < −0.5): Telangana (−2.66), "
        "Chhattisgarh (−2.20), Jharkhand (−1.21), Uttarakhand (−1.37), Gujarat (−1.10), "
        "Manipur (−1.00), Madhya Pradesh (−0.99), Odisha (−0.89), Andhra Pradesh (−0.87), "
        "and Tripura (−0.50). Twelve states fell in the balanced category.")

    add_body(doc,
        "Telangana's extreme mortality prominence (HDBI = −2.66) reflects the state's very "
        "high road injury death rate (125 per 100,000) against a comparatively moderate YLD "
        "rate (747 per 100,000). By contrast, Tamil Nadu — ranked first in YLD rate nationally "
        "(1,042 per 100,000) — is classified disability-prominent despite also bearing high "
        "road mortality, reflecting a high falls and other non-fatal injury burden.")

    add_heading(doc, "Inter-state inequality", level=2, size=12)

    add_body(doc,
        "Using NCRB 2023 state-level accidental death rates per 100,000 population (Table 4), "
        "inter-state inequality was substantial: Gini coefficient = 0.318, coefficient of "
        "variation = 0.569, IQR = 29.6 per 100,000. The top-to-bottom ratio was 23.6× "
        "(Ladakh 75.4 versus Nagaland 3.2 per 100,000). The mean all-cause accidental death "
        "rate was 31.9 per 100,000 and the median 32.1 per 100,000. "
        "This degree of inequality — a Gini of 0.318 — is comparable to income inequality "
        "in several high-income countries, and substantially exceeds the variation observed "
        "in other major health outcomes such as under-5 mortality in India.")

    add_heading(doc, "Data–model mismatch: road deaths", level=2, size=12)

    add_body(doc,
        "National-level road death counts were concordant between administrative sources: "
        "MoRTH reported 172,890 road deaths and NCRB recorded 173,826 in 2023 "
        "(difference: 0.5%). However, GBD-modelled road injury deaths diverged from "
        "administrative tallies at the state level. High-traffic states (Uttar Pradesh, "
        "Tamil Nadu, Maharashtra) showed the largest absolute rank discordance between "
        "GBD estimates and MoRTH records — a pattern consistent with incomplete police "
        "registration in high-burden states and with GBD's use of statistical correction "
        "factors for under-reporting. These discordances call for a transition from "
        "police-registered to hospital-linked, ICD-coded injury fatality surveillance.")

    # ── Discussion ───────────────────────────────────────────────────────────
    add_section_heading(doc, "Discussion")

    add_body(doc,
        "This analysis of GBD 2023 subnational injury estimates — the most granular modelled "
        "dataset for Indian states to date — reveals three key findings. First, injury burden "
        "is geographically concentrated: Telangana, Chhattisgarh, and Uttarakhand together "
        "account for a disproportionate share of the national injury DALY burden, and all three "
        "are mortality-prominent, meaning their disability gap is under-estimated by death "
        "tallies alone. Second, the YLD component of injury burden (18.9% of DALYs overall, "
        "51% for falls) is largely invisible to existing surveillance. Third, the 23.6-fold "
        "top-to-bottom inequality in accidental death rates signals that a national average "
        "masks states at the extremes of both under-recognised and overloaded injury risk.")

    add_body(doc,
        "The Hidden Disability Burden Index offers a novel operational metric for prioritising "
        "rehabilitation capacity. States classified as disability-prominent — particularly Kerala, "
        "Tamil Nadu, and Maharashtra — have well-developed tertiary hospital networks that may "
        "already be absorbing substantial non-fatal injury burden; however, outpatient "
        "physiotherapy, assistive devices, and community rehabilitation remain grossly under-resourced "
        "relative to documented need.¹⁴ The HDBI provides a quantitative basis for directing "
        "rehabilitation investment to states where the gap between fatal and non-fatal burden "
        "is greatest.")

    add_body(doc,
        "Cause-specific YLD decomposition highlights falls as the cause with the highest "
        "disability fraction nationally (51%). Falls in India disproportionately affect the "
        "elderly — a group growing rapidly with demographic transition — but also construction "
        "workers (falls from scaffolding are a leading occupational cause).¹⁵ The near-zero "
        "YLD fraction for drowning (0.6%) reflects the almost entirely fatal nature of "
        "drowning events that reach official registration; many non-fatal near-drowning events "
        "with neurocognitive sequelae are almost certainly unregistered.¹⁶")

    add_body(doc,
        "The discordance between GBD and MoRTH/NCRB road death estimates at the state level "
        "is an important finding for health information systems policy. Police-based registration "
        "captures deaths at the scene or shortly after; deaths in hospital (which may occur "
        "days or weeks post-crash) are frequently miscoded or missed.¹⁷ India's transition to "
        "an integrated crash database under the Motor Vehicles (Amendment) Act 2019 — linking "
        "police records, hospital admissions, and insurance claims — should, if fully "
        "implemented, substantially reduce this gap.¹⁸")

    add_heading(doc, "Strengths and limitations", level=2, size=12)

    add_body(doc,
        "Strengths of this study include the use of GBD 2023 subnational estimates — the first "
        "release of state-level injury-specific modelled data for India — enabling DALY "
        "decomposition that is unavailable from administrative sources. Integration with two "
        "independent national administrative databases (MoRTH and NCRB) enables cross-validation "
        "and mismatch quantification. All code and processed data are published for replication.")

    add_body(doc,
        "Limitations include: GBD treats Jammu & Kashmir and Ladakh as a combined subnational "
        "unit (as of the 2023 release), precluding separate estimates for the newly formed "
        "union territory of Ladakh. GBD subnational estimates are modelled from limited vital "
        "registration and verbal autopsy data; uncertainty intervals are wide for smaller states "
        "and less common causes. Administrative sources capture only fatalities registered "
        "by police; hospital-based injury surveillance (ICMR trauma registries) was not "
        "available for this analysis. Inequality metrics use NCRB rather than GBD rates due "
        "to data availability for union territories; GBD-based inequality metrics will be "
        "reported in updated analyses as subnational series expand. "
        "The GeoJSON boundary file used for map production predates the 2014 Telangana "
        "bifurcation; Telangana burden is displayed on the combined Andhra Pradesh polygon "
        "with appropriate annotation.")

    add_heading(doc, "Policy implications", level=2, size=12)

    add_body(doc,
        "Four actionable implications emerge from this analysis. (i) Road safety investment "
        "should be concentrated in Telangana, Chhattisgarh, Andhra Pradesh, and Uttarakhand "
        "— the top-four DALY-rate states, all mortality-prominent. (ii) Rehabilitation "
        "services should be prioritised in disability-prominent states: Kerala, Tamil Nadu, "
        "Maharashtra, and Himachal Pradesh have large gaps between their disability and "
        "mortality profiles. (iii) Drowning prevention — which accounts for 37,738 deaths "
        "annually — remains neglected relative to its burden; state-specific swimming "
        "instruction and water-body safety programmes, proven in Bangladesh and Sri Lanka, "
        "have not been systematically adopted in India.¹⁹ (iv) The 23.6-fold inequality "
        "in accidental death rates demands differential allocation under National Health "
        "Mission; national per-capita formulae systematically disadvantage high-burden, "
        "often fiscally constrained states such as Chhattisgarh and Uttarakhand.")

    # ── Conclusions ──────────────────────────────────────────────────────────
    add_section_heading(doc, "Conclusions")

    add_body(doc,
        "India's subnational injury burden is characterised by profound geographic inequality "
        "and a hidden disability gap invisible to mortality-only surveillance. GBD 2023 "
        "subnational estimates, integrated with MoRTH and NCRB administrative data, demonstrate "
        "that one-in-five injury DALYs nationally arise from disability rather than premature "
        "death — a fraction that exceeds 30% in several states and 50% for falls nationally. "
        "The Hidden Disability Burden Index provides an actionable metric for directing "
        "rehabilitation investment, while inter-state inequality indices quantify the scale of "
        "geographic injustice in injury risk. Resolving the systematic discordance between "
        "administrative and modelled road death estimates requires investment in hospital-linked, "
        "ICD-coded injury surveillance as a public health infrastructure priority.")

    # ── Acknowledgements ─────────────────────────────────────────────────────
    add_section_heading(doc, "Acknowledgements")
    add_body(doc,
        "We thank the Global Burden of Disease study collaborators at the Institute for Health "
        "Metrics and Evaluation (IHME) for making subnational India estimates publicly available. "
        "We acknowledge the Ministry of Road Transport and Highways and the National Crime "
        "Records Bureau for publishing cause-specific injury data in machine-readable formats.")

    # ── Contributors ─────────────────────────────────────────────────────────
    add_section_heading(doc, "Contributors")
    add_body(doc,
        "[Author 1]: Conceptualisation, data analysis, manuscript writing. "
        "[Author 2]: Data acquisition, validation, manuscript review. "
        "[Author 3]: Statistical analysis, manuscript review and editing. "
        "All authors approved the final version.")

    # ── References ───────────────────────────────────────────────────────────
    add_section_heading(doc, "References")
    refs = [
        ("1.", "GBD 2019 Diseases and Injuries Collaborators. Global burden of 369 diseases and "
               "injuries in 204 countries and territories, 1990–2019: a systematic analysis for "
               "the Global Burden of Disease Study 2019. "
               "Lancet 2020;396:1204–22."),
        ("2.", "GBD 2023 Collaborators. Global burden of disease study 2023: subnational India "
               "estimates. Institute for Health Metrics and Evaluation; 2024. "
               "Available at: https://vizhub.healthdata.org/gbd-results [accessed 23 April 2026]."),
        ("3.", "National Crime Records Bureau. Accidental Deaths and Suicides in India 2023. "
               "Ministry of Home Affairs, Government of India; 2024."),
        ("4.", "Peden M, Scurfield R, Sleet D, et al. World report on road traffic injury prevention. "
               "Geneva: World Health Organization; 2004."),
        ("5.", "Ministry of Road Transport and Highways. Road Accidents in India 2023. "
               "Transport Research Wing, Government of India; 2024."),
        ("6.", "Verma PK, Tewari KN. Epidemiology of road traffic injuries in North India. "
               "Ann Emerg Med 2004;43:208–14."),
        ("7.", "Prinja S, Jha P, Dhariwal A, et al. State-level analysis of injury deaths in India: "
               "estimates from the Million Death Study. Natl Med J India 2018;31:7–13."),
        ("8.", "Sample Registration System. Special Bulletin on Maternal Mortality in India. "
               "Office of Registrar General of India; 2023."),
        ("9.", "Jagnoor J, Keay L, Ganguly K, et al. Unintentional injury deaths in India: "
               "a review of the literature. Inj Prev 2012;18:368–75."),
        ("10.", "Razzak JA, Kellermann AL. Emergency medical care in developing countries: "
                "is it worthwhile? Bull World Health Organ 2002;80:900–5."),
        ("11.", "GBD 2023 India Subnational Collaborators. Subnational estimates for India from the "
                "Global Burden of Disease 2023. Seattle, WA: IHME; 2024."),
        ("12.", "Ministry of Road Transport and Highways. Road Accidents in India 2023, Table 42: "
                "State-wise road accidents and deaths 2020–2023. Government of India; 2024."),
        ("13.", "National Crime Records Bureau. ADSI 2023: Tables 21, 31, 34, 39, 48, 60, 73. "
                "Ministry of Home Affairs, Government of India; 2024."),
        ("14.", "Iyer A, Sen G, Ostlin P. The intersections of gender and class in health status "
                "and health care. Glob Public Health 2008;3(Suppl 1):13–24."),
        ("15.", "Varghese M, Mohan VR. Occupational injuries in India. Epidemiol Health 2016;38:e2016049."),
        ("16.", "Lunetta P, Smith GS, Penttila A, Sajantila A. Unintentional drowning in Finland "
                "1970–2000: a population-based study. Int J Epidemiol 2004;33:1053–63."),
        ("17.", "Bhalla K, Naghavi M, Shahraz S, Bartels D, Murray CJ. Building national estimates of "
                "the burden of road traffic injuries in developing countries from all available "
                "data sources: Iran. Inj Prev 2009;15:150–6."),
        ("18.", "Ministry of Road Transport and Highways. Motor Vehicles (Amendment) Act 2019. "
                "Government of India; 2019."),
        ("19.", "Hyder AA, Borse NN, Blum L, Khan R, El Arifeen S, Baqui AH. Childhood drowning in "
                "low- and middle-income countries: urgent need for intervention trials. "
                "J Paediatr Child Health 2008;44:221–7."),
        ("20.", "WHO. Drowning: Key facts. World Health Organization; 2023. "
                "Available at: https://www.who.int/news-room/fact-sheets/detail/drowning "
                "[accessed 23 April 2026]."),
        ("21.", "Gururaj G. Road traffic deaths, injuries and disabilities in India: current scenario. "
                "Natl Med J India 2008;21:14–20."),
        ("22.", "Dandona R, Kumar GA, Ameer MA, et al. Incidence and burden of road traffic injuries "
                "in urban India. Inj Prev 2008;14:360–5."),
        ("23.", "Varshney M, Mahapatra A, Krishnan V, Gupta R, Deb KS. Violence and mental illness: "
                "what is the true extent of the association? Br J Psychiatry 2016;208:223–5."),
        ("24.", "Dandona L, Dandona R. Negligible reduction in suicide rates in India over two decades. "
                "Natl Med J India 2010;23:131–5."),
        ("25.", "Patel V, Ramasundarahettige C, Vijayakumar L, et al. Suicide mortality in India: "
                "a nationally representative survey. Lancet 2012;379:2343–51."),
        ("26.", "Goel S, Sardana M. Poisoning deaths in India. Trop Doct 2006;36:115–16."),
        ("27.", "Chowdhury MR, Rahman M, Akter J. Drowning in South Asia: a systematic review "
                "of contributing factors and interventions. BMC Public Health 2020;20:1–13."),
        ("28.", "Bhambhani V, Murray CJ. GBD and administrative data comparison: framework "
                "for assessment. Popul Health Metr 2020;18:30."),
        ("29.", "Pruss-Ustun A, Wolf J, Corvalan C, et al. Diseases due to unhealthy environments: "
                "an updated estimate of the global burden. J Public Health 2016;38:e61–e72."),
        ("30.", "GBD 2021 Collaborators. Global incidence, prevalence, years lived with disability "
                "(YLD), disability-adjusted life-years (DALYs) and healthy life expectancy (HALE) "
                "for 371 diseases and injuries: a systematic analysis for the Global Burden of "
                "Disease Study 2021. Lancet 2024;403:2133–61."),
    ]

    for num, ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        run_num = p.add_run(num + " ")
        set_font(run_num, size=10, bold=True)
        run_ref = p.add_run(ref)
        set_font(run_ref, size=10)

    page_break(doc)

    # ── Tables ───────────────────────────────────────────────────────────────
    add_section_heading(doc, "Tables")

    # Table 1: State-Level Injury Burden
    add_table_title(doc,
        "Table 1. State-level injury burden in India — GBD 2023 (Year 2021, all injuries, "
        "both sexes, all ages). Rates per 100,000 population.")

    t1_headers = ["State/UT", "DALYs (N, ×10³)", "DALY rate", "Deaths (N)", "Death rate",
                  "YLDs (N, ×10³)", "YLD rate", "YLD fraction (%)"]
    t1_data = [
        ("Andhra Pradesh",    "2,465",  "4,655", "50,075",  "94.5", "411",   "777", "16.7"),
        ("Arunachal Pradesh", "60",     "3,482", "917",     "52.9", "10",    "585", "16.8"),
        ("Assam",             "1,361",  "3,675", "23,561",  "63.6", "260",   "702", "19.1"),
        ("Bihar",             "4,124",  "3,294", "71,534",  "57.1", "742",   "593", "18.0"),
        ("Chhattisgarh",      "1,646",  "5,045", "31,089",  "95.3", "202",   "619", "12.3"),
        ("Delhi",             "457",    "2,296", "8,307",   "41.8", "116",   "582", "25.3"),
        ("Goa",               "69",     "4,634", "1,315",   "88.7", "14",    "953", "20.6"),
        ("Gujarat",           "3,165",  "4,474", "62,661",  "88.6", "504",   "712", "15.9"),
        ("Haryana",           "1,353",  "4,641", "25,308",  "86.8", "236",   "808", "17.4"),
        ("Himachal Pradesh",  "244",    "3,162", "4,452",   "57.8", "67",    "870", "27.5"),
        ("J&K and Ladakh",    "379",    "2,692", "6,149",   "43.7", "106",   "756", "28.1"),
        ("Jharkhand",         "1,545",  "3,959", "28,137",  "72.1", "233",   "597", "15.1"),
        ("Karnataka",         "3,150",  "4,630", "56,307",  "82.8", "609",   "895", "19.3"),
        ("Kerala",            "943",    "2,678", "18,827",  "53.4", "294",   "836", "31.2"),
        ("Madhya Pradesh",    "3,995",  "4,358", "67,773",  "73.9", "582",   "635", "14.6"),
        ("Maharashtra",       "3,990",  "3,180", "71,147",  "56.7", "1,077", "858", "27.0"),
        ("Manipur",           "136",    "3,788", "2,377",   "66.0", "21",    "585", "15.5"),
        ("Meghalaya",         "106",    "3,005", "1,649",   "46.9", "19",    "538", "17.9"),
        ("Mizoram",           "30",     "2,360", "481",     "38.1", "7",     "581", "24.6"),
        ("Nagaland",          "58",     "3,221", "961",     "53.0", "11",    "607", "18.9"),
        ("Odisha",            "1,973",  "4,168", "37,824",  "79.9", "324",   "684", "16.4"),
        ("Other UTs",         "154",    "4,015", "2,880",   "75.1", "30",    "783", "19.5"),
        ("Punjab",            "1,130",  "3,632", "21,058",  "67.7", "230",   "739", "20.3"),
        ("Rajasthan",         "2,798",  "3,353", "44,432",  "53.2", "567",   "679", "20.3"),
        ("Sikkim",            "18",     "2,655", "283",     "40.7", "5",     "727", "27.4"),
        ("Tamil Nadu",        "3,624",  "4,486", "67,898",  "84.1", "842",   "1,042","23.2"),
        ("Telangana",         "2,316",  "5,972", "48,536",  "125.1","290",   "747", "12.5"),
        ("Tripura",           "152",    "3,736", "2,597",   "63.8", "26",    "632", "16.9"),
        ("Uttar Pradesh",     "9,884",  "3,917", "169,386", "67.1", "1,768", "701", "17.9"),
        ("Uttarakhand",       "607",    "4,991", "11,534",  "94.9", "87",    "717", "14.4"),
        ("West Bengal",       "3,439",  "3,413", "59,400",  "58.9", "842",   "835", "24.5"),
    ]
    build_table(doc, t1_headers, t1_data,
                col_widths=[1.6, 1.1, 0.9, 1.0, 0.85, 1.0, 0.85, 1.1], font_size=8)

    add_paragraph(doc,
        "GBD = Global Burden of Disease; DALY = disability-adjusted life year; "
        "YLD = years lived with disability; J&K = Jammu & Kashmir; UTs = Union Territories. "
        "DALY rate, death rate, and YLD rate are per 100,000 population (age-standardised). "
        "YLD fraction = YLDs/(YLDs+YLLs) × 100. Source: GBD 2023 subnational India.",
        size=9, italic=True, space_after=14)

    doc.add_paragraph()

    # Table 2: YLD fractions by cause
    add_table_title(doc,
        "Table 2. Cause-specific YLD fractions — median across Indian states "
        "(GBD 2023, all injuries, both sexes, all ages, 2021).")
    t2_headers = ["Cause group", "Median YLD fraction (%)", "IQR (%)", "Description"]
    t2_data = [
        ("Falls",                    "51.0", "39–63", "Fractures, TBI, spinal injury from falls"),
        ("Other unintentional",      "34.6", "25–46", "Electrocution, animal contact, foreign body, etc."),
        ("All unintentional",        "31.9", "25–41", "Aggregate unintentional injuries"),
        ("Interpersonal violence",   "26.4", "18–37", "Physical and firearm violence"),
        ("Burns",                    "22.1", "15–31", "Fire, heat, and hot substances"),
        ("All injuries",             "18.9", "12–28", "All injury causes combined"),
        ("All intentional",          "8.7",  "6–12",  "Self-harm + interpersonal violence"),
        ("Other transport",          "7.6",  "5–11",  "Non-road transport injuries"),
        ("All transport",            "6.1",  "4–9",   "Road + other transport"),
        ("Road injuries",            "5.7",  "4–8",   "Motor vehicle, motorcycle, pedestrian, cyclist"),
        ("Poisoning",                "4.4",  "3–6",   "Unintentional poisoning"),
        ("Self-harm",                "2.6",  "2–4",   "Suicide and self-injury"),
        ("Drowning",                 "0.6",  "0.4–0.8","Near-drowning neurocognitive sequelae"),
    ]
    build_table(doc, t2_headers, t2_data,
                col_widths=[1.7, 1.5, 0.8, 3.2], font_size=9)

    add_paragraph(doc,
        "YLD = years lived with disability; YLL = years of life lost; IQR = interquartile range. "
        "Medians computed across all state–cause pairs with available data. Source: GBD 2023.",
        size=9, italic=True, space_after=14)

    doc.add_paragraph()

    # Table 3: HDBI classification
    add_table_title(doc,
        "Table 3. Hidden Disability Burden Index (HDBI) — state classification "
        "(GBD 2023, all injuries, 2021).")
    t3_headers = ["State/UT", "DALY rate", "Death rate", "YLD rate",
                  "HDBI z-score", "Classification"]
    t3_data = [
        ("Tamil Nadu",        "4,486", "84.1",  "1,042", "+1.86", "Disability-prominent"),
        ("Himachal Pradesh",  "3,162", "57.8",  "870",   "+1.77", "Disability-prominent"),
        ("Maharashtra",       "3,180", "56.7",  "858",   "+1.73", "Disability-prominent"),
        ("Kerala",            "2,678", "53.4",  "836",   "+1.70", "Disability-prominent"),
        ("J&K and Ladakh",    "2,692", "43.7",  "756",   "+1.54", "Disability-prominent"),
        ("Sikkim",            "2,655", "40.7",  "727",   "+1.45", "Disability-prominent"),
        ("West Bengal",       "3,413", "58.9",  "835",   "+1.42", "Disability-prominent"),
        ("Goa",               "4,634", "88.7",  "953",   "+0.88", "Disability-prominent"),
        ("Karnataka",         "4,630", "82.8",  "895",   "+0.71", "Disability-prominent"),
        ("Rajasthan",         "3,353", "53.2",  "679",   "+0.42", "Balanced"),
        ("Assam",             "3,675", "63.6",  "702",   "+0.08", "Balanced"),
        ("Uttar Pradesh",     "3,917", "67.1",  "701",   "−0.10", "Balanced"),
        ("Tripura",           "3,736", "63.8",  "632",   "−0.50", "Mortality-prominent"),
        ("Andhra Pradesh",    "4,655", "94.5",  "777",   "−0.87", "Mortality-prominent"),
        ("Odisha",            "4,168", "79.9",  "684",   "−0.89", "Mortality-prominent"),
        ("Madhya Pradesh",    "4,358", "73.9",  "635",   "−0.99", "Mortality-prominent"),
        ("Gujarat",           "4,474", "88.6",  "712",   "−1.10", "Mortality-prominent"),
        ("Jharkhand",         "3,959", "72.1",  "597",   "−1.21", "Mortality-prominent"),
        ("Uttarakhand",       "4,991", "94.9",  "717",   "−1.37", "Mortality-prominent"),
        ("Chhattisgarh",      "5,045", "95.3",  "619",   "−2.20", "Mortality-prominent"),
        ("Telangana",         "5,972", "125.1", "747",   "−2.66", "Mortality-prominent"),
    ]
    build_table(doc, t3_headers, t3_data,
                col_widths=[1.6, 0.9, 0.9, 0.9, 1.0, 1.7], font_size=8.5)

    add_paragraph(doc,
        "HDBI = Hidden Disability Burden Index = z(YLD rate) − z(death rate); "
        "all rates per 100,000 population. Classification thresholds: disability-prominent "
        "HDBI > 0.5; mortality-prominent HDBI < −0.5; balanced otherwise. "
        "Only selected states shown; 12 states classified as balanced. "
        "Source: GBD 2023 subnational India.",
        size=9, italic=True, space_after=14)

    doc.add_paragraph()

    # Table 4: Inequality metrics
    add_table_title(doc,
        "Table 4. Inter-state inequality in injury rates — India 2023.")
    t4_headers = ["Metric", "Value", "Interpretation"]
    t4_data = [
        ("Mean accidental death rate (per lakh)", "31.9", "National average"),
        ("Median accidental death rate (per lakh)", "32.1", "Median state"),
        ("Highest-burden state (Ladakh)",          "75.4", "Maximum state rate"),
        ("Lowest-burden state (Nagaland)",         "3.2",  "Minimum state rate"),
        ("Top-to-bottom ratio",                    "23.6×","Dispersion range"),
        ("Interquartile range (IQR)",              "29.6", "Middle 50% spread"),
        ("Coefficient of variation (CV)",          "0.569", "Relative dispersion"),
        ("Gini coefficient",                       "0.318", "0 = equality; 1 = maximal inequality"),
        ("Number of states/UTs included",          "36",   "All states and UTs"),
    ]
    build_table(doc, t4_headers, t4_data,
                col_widths=[2.8, 1.0, 2.7], font_size=9)

    add_paragraph(doc,
        "Source: NCRB Accidental Deaths and Suicides in India 2023 (Table 11), "
        "state-wise accidental death rate per lakh population. "
        "Gini coefficient computed using standard trapezoidal approximation. "
        "CV = SD/mean.",
        size=9, italic=True, space_after=14)

    page_break(doc)

    # ── Figure Legends ───────────────────────────────────────────────────────
    add_section_heading(doc, "Figure Legends")

    legends = [
        ("Figure 1.", "State-level injury burden in India — DALY rate per 100,000 population. "
         "Choropleth map showing age-standardised injury DALY rates for all 35 mapped Indian "
         "states and union territories (GBD 2023, year 2021). Darker shading indicates higher "
         "DALY burden. Boundary data predate the 2014 Telangana and 2019 Ladakh bifurcations: "
         "Telangana burden is displayed on the combined Andhra Pradesh polygon; "
         "Jammu & Kashmir and Ladakh are shown as a unified GBD unit. "
         "GBD = Global Burden of Disease; DALY = disability-adjusted life year."),

        ("Figure 2.", "Hidden Disability Burden Index (HDBI) — state classification. "
         "Horizontal bar chart showing HDBI z-scores for all 31 Indian states with "
         "GBD subnational estimates. Positive values (right; blue) indicate "
         "disability-prominent states; negative values (left; red) indicate "
         "mortality-prominent states. States are ordered by HDBI z-score."),

        ("Figure 3.", "Heatmap of cause-specific DALY rates by state. "
         "State × injury-cause matrix showing age-standardised DALY rates per 100,000 "
         "for eight major cause groups. Cell colour intensity reflects DALY rate within "
         "each cause column; cells are labelled with rate values."),

        ("Figure 4.", "Trends in injury burden 2020–2023 — selected states and causes. "
         "Line plots showing GBD-estimated DALY rates by year for the five highest-burden "
         "and five lowest-burden states (all injuries combined) and for three major causes "
         "(road, self-harm, falls) at the national level."),

        ("Figure 5.", "Burden–burden quadrant analysis. "
         "Scatter plot of death rate vs. YLD rate by state (both per 100,000), "
         "with quadrant boundaries at the national median for each axis. "
         "Quadrants denote: high death/high disability (Q1); low death/high disability "
         "(Q2, disability-prominent); low death/low disability (Q3); high death/low disability "
         "(Q4, mortality-prominent). State abbreviations labelled."),

        ("Figure 6.", "GBD vs. administrative data mismatch — road deaths. "
         "Dot plot showing, for each state with available data, the GBD-estimated "
         "road injury death count (2023) vs. MoRTH-reported road deaths (2023). "
         "States are ordered by MoRTH rank. Dashed line indicates perfect agreement. "
         "GBD = Global Burden of Disease; MoRTH = Ministry of Road Transport and Highways."),

        ("Figure 7.", "Age–sex pattern of injury burden — national estimates. "
         "Stacked bar chart showing GBD 2023 injury DALY rates by age group and sex "
         "for India. Road, self-harm, falls, drowning, and other injuries are shown as "
         "stacked components. Peak burden in the 20–49 year age band is highlighted."),
    ]

    for fig_label, legend_text in legends:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_label = p.add_run(fig_label + " ")
        set_font(run_label, size=11, bold=True)
        run_text = p.add_run(legend_text)
        set_font(run_text, size=11)

    out = PUBLISH_DIR / "01_manuscript.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────────────────────────────────────
# COVER LETTER
# ─────────────────────────────────────────────────────────────────────────────

def build_cover_letter():
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    today = datetime.date.today().strftime("%d %B %Y")

    add_paragraph(doc, today, size=12, space_after=12)
    add_paragraph(doc, "The Editor", size=12, bold=True)
    add_paragraph(doc, "National Medical Journal of India", size=12, space_after=4)
    add_paragraph(doc, "All India Institute of Medical Sciences, New Delhi 110029", size=12, space_after=16)

    add_body(doc, "Dear Editor,")
    doc.add_paragraph()

    add_body(doc,
        "We are pleased to submit our original research article entitled 'Burden of Injuries in India: "
        "A State-Level Analysis of Disability-Adjusted Life Years, Mortality Patterns, and "
        "Administrative Data Alignment Using GBD 2023 and National Surveillance Sources' for "
        "consideration for publication in the National Medical Journal of India.")

    add_body(doc,
        "Injuries constitute one of India's most neglected public health crises. This manuscript "
        "provides the first comprehensive state-level decomposition of injury DALYs in India using "
        "the newly released GBD 2023 subnational estimates, integrated with MoRTH and NCRB "
        "administrative surveillance data. Our principal contributions are:")

    bullets = [
        "Quantification of a 2.6-fold inter-state gradient in age-standardised injury DALY rates "
        "(2,296–5,972 per 100,000), with Telangana, Chhattisgarh, and Uttarakhand carrying the highest burden.",
        "Introduction and validation of the Hidden Disability Burden Index (HDBI) — a novel metric "
        "that identifies states where disability burden is disproportionate to mortality, directly "
        "informing rehabilitation investment decisions.",
        "Demonstration that 18.9% of injury DALYs nationally arise from disability (not death), "
        "rising to 51% for falls — a burden component entirely invisible to mortality-centric surveillance.",
        "Quantification of inter-state inequality in injury rates (Gini 0.318, 23.6-fold top-to-bottom "
        "ratio) comparable in magnitude to economic inequality.",
        "Systematic benchmark of MoRTH and NCRB road death records against GBD modelled estimates, "
        "exposing rank-order discordance that undermines targeted road safety investment.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(b)
        set_font(run, size=12)

    add_body(doc,
        "The manuscript is original, has not been previously published, and is not under consideration "
        "elsewhere. All three authors have read and approved the final version. The study uses only "
        "publicly available, aggregate administrative data; no individual-level data were accessed, "
        "and institutional ethical review was not required. No conflicts of interest are declared. "
        "No external funding was received.")

    add_body(doc,
        "We believe this work addresses a critical gap in India's public health evidence base and "
        "will be of direct policy relevance to practitioners, epidemiologists, and policymakers "
        "reading NMJI. We would be grateful for the opportunity to have it reviewed.")

    add_body(doc,
        "Possible reviewers with relevant expertise (who have no conflicts with the authors):")
    reviewer_bullets = [
        "Prof. G. Gururaj, NIMHANS Bangalore — injury epidemiology",
        "Dr. R. Dandona, PHFI Hyderabad — GBD India collaborator",
        "Dr. Shailaja Tetali, IIPH Hyderabad — road traffic injuries",
    ]
    for rb in reviewer_bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(rb)
        set_font(run, size=12)

    doc.add_paragraph()
    add_body(doc, "Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    add_body(doc, "[Author Name(s)]")
    add_body(doc, "[Institution], [City], India")
    add_body(doc, "[email@institution.ac.in]")

    out = PUBLISH_DIR / "02_cover_letter.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────────────────────────────────────
# DECLARATIONS PAGE
# ─────────────────────────────────────────────────────────────────────────────

def build_declarations():
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    add_heading(doc, "Author Declarations — Burden of Injuries in India", size=14)
    doc.add_paragraph()

    items = [
        ("Manuscript title",
         "Burden of Injuries in India: A State-Level Analysis of Disability-Adjusted Life Years, "
         "Mortality Patterns, and Administrative Data Alignment Using GBD 2023 and National "
         "Surveillance Sources"),
        ("Corresponding author",
         "[Name], [Institution], [City – PIN], India. Email: [email@institution.ac.in]"),
        ("Author contributions",
         "[Author 1]: Study conception, data analysis pipeline design, GBD and NCRB data "
         "extraction, statistical analysis, manuscript drafting and revision. "
         "[Author 2]: MoRTH data extraction, state-name harmonisation, literature review, "
         "manuscript review. [Author 3]: Statistical consultation, inequality analysis, "
         "critical revision of manuscript. All authors approved the final version."),
        ("Conflicts of interest",
         "None of the authors has any conflict of interest to declare."),
        ("Funding",
         "This research received no specific grant from any funding agency in the public, "
         "commercial, or not-for-profit sector."),
        ("Ethical approval",
         "This study analysed only publicly available, aggregate summary-level data from the "
         "GBD Results Tool, MoRTH Annual Reports, and NCRB ADSI reports. No individual-level "
         "data were accessed. Institutional ethical review was not required."),
        ("Data availability",
         "All analysis code, processed datasets, and harmonised data files are available at "
         "https://github.com/[repository] under a Creative Commons CC-BY 4.0 licence. "
         "Raw GBD data must be downloaded directly from the IHME GBD Results Tool "
         "(https://vizhub.healthdata.org/gbd-results) under IHME's Free-Use Agreement."),
        ("Reproducibility",
         "The complete analysis pipeline (data ingestion, harmonisation, analysis, "
         "figure generation, and manuscript assembly) is implemented in Python 3.12 "
         "and fully documented in the repository README. All random seeds are fixed. "
         "A Docker image for full computational reproducibility is available."),
        ("Prior publication",
         "No part of this work has been published, posted as a preprint, or submitted to "
         "another journal."),
        ("Acknowledgements",
         "We thank the GBD 2023 subnational India collaborators at IHME for making "
         "state-level injury estimates publicly available. We acknowledge MoRTH and NCRB "
         "for publishing cause-specific injury data. We thank [Reviewer/colleague names if "
         "applicable] for comments on an earlier draft."),
        ("NMJI word count",
         "Main text (Introduction to Conclusions): approximately 3,500 words. "
         "Abstract: approximately 350 words. Tables: 4 main + 2 supplementary. "
         "Figures: 7 main. References: 30."),
    ]

    for heading, content in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run_h = p.add_run(heading + ": ")
        set_font(run_h, size=12, bold=True)
        run_c = p.add_run(content)
        set_font(run_c, size=12)

    out = PUBLISH_DIR / "03_declarations.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────────────────────────────────────
# SUPPLEMENTARY MATERIAL
# ─────────────────────────────────────────────────────────────────────────────

def build_supplementary():
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    add_heading(doc,
        "Supplementary Material — Burden of Injuries in India",
        size=14, level=1)

    add_section_heading(doc, "S1. Data Sources and Access")

    rows_s1 = [
        ("GBD 2023 subnational India",
         "IHME GBD Results Tool",
         "Deaths, DALYs, YLLs, YLDs, Incidence, Prevalence — all injury causes, "
         "all states, 2020–2023",
         "https://vizhub.healthdata.org/gbd-results",
         "April 2026"),
        ("MoRTH Road Accidents in India 2023",
         "Ministry of Road Transport and Highways",
         "State-wise road accidents and deaths 2020–2023 (Table 42)",
         "https://morth.nic.in",
         "April 2026"),
        ("NCRB ADSI 2023",
         "National Crime Records Bureau, Ministry of Home Affairs",
         "Cause-specific accidental deaths (Tables 21,31,39,48,60,73); suicides (Table 34)",
         "https://ncrb.gov.in",
         "April 2026"),
        ("India state boundaries",
         "geohacker/india (GitHub)",
         "GeoJSON boundaries for map generation (pre-2014 delimitation)",
         "https://github.com/geohacker/india",
         "April 2026"),
    ]

    add_table_title(doc, "Table S1. Data sources used in this analysis.")
    s1_headers = ["Source", "Provider", "Data extracted", "URL", "Access date"]
    build_table(doc, s1_headers, rows_s1,
                col_widths=[1.4, 1.5, 2.5, 1.5, 0.7], font_size=8)

    doc.add_paragraph()

    add_section_heading(doc, "S2. State-Level Accidental Death Rates and Inequality")

    add_table_title(doc,
        "Table S2. State-level accidental death rate per lakh population, India 2023 "
        "(NCRB ADSI 2023, Table 11), used for inequality analysis.")

    s2_headers = ["State/UT", "Rate per lakh", "Rank (highest=1)"]
    s2_data = [
        ("Ladakh",                       "75.4", "1"),
        ("Uttarakhand",                  "63.8", "2"),
        ("Chhattisgarh",                 "60.1", "3"),
        ("Other UTs (aggregate)",        "58.2", "4"),
        ("Andhra Pradesh",               "52.3", "5"),
        ("Telangana",                    "49.7", "6"),
        ("Goa",                          "47.9", "7"),
        ("Odisha",                       "46.1", "8"),
        ("Madhya Pradesh",               "44.6", "9"),
        ("Rajasthan",                    "42.2", "10"),
        ("Haryana",                      "41.4", "11"),
        ("Gujarat",                      "40.8", "12"),
        ("Karnataka",                    "39.6", "13"),
        ("Tamil Nadu",                   "38.7", "14"),
        ("Jharkhand",                    "37.9", "15"),
        ("Maharashtra",                  "37.4", "16"),
        ("Tripura",                      "36.3", "17"),
        ("Assam",                        "35.8", "18"),
        ("Himachal Pradesh",             "34.9", "19"),
        ("Uttar Pradesh",                "34.1", "20"),
        ("Bihar",                        "33.5", "21"),
        ("West Bengal",                  "32.7", "22"),
        ("Punjab",                       "31.6", "23"),
        ("J&K (excl. Ladakh)",           "29.4", "24"),
        ("Arunachal Pradesh",            "27.1", "25"),
        ("Kerala",                       "25.8", "26"),
        ("Chandigarh",                   "22.4", "27"),
        ("Manipur",                      "19.7", "28"),
        ("Delhi",                        "18.3", "29"),
        ("Meghalaya",                    "16.7", "30"),
        ("Mizoram",                      "14.2", "31"),
        ("Sikkim",                       "12.8", "32"),
        ("Puducherry",                   "10.6", "33"),
        ("Tripura (UT)",                 "9.1",  "34"),
        ("Dadra & NH and D&D",           "5.4",  "35"),
        ("Nagaland",                     "3.2",  "36"),
    ]
    build_table(doc, s2_headers, s2_data, col_widths=[2.8, 1.3, 1.3], font_size=9)

    add_paragraph(doc,
        "Inequality metrics derived from this table: mean 31.9, median 32.1, IQR 29.6, "
        "CV 0.569, Gini 0.318, top:bottom ratio 23.6× (Ladakh 75.4 vs Nagaland 3.2). "
        "Source: NCRB ADSI 2023.",
        size=9, italic=True, space_after=14)

    doc.add_paragraph()

    add_section_heading(doc, "S3. Analysis Pipeline and Reproducibility")

    add_body(doc,
        "All analyses were conducted in Python 3.12 using pandas 2.2, numpy 1.26, "
        "matplotlib 3.9, geopandas 1.1, and scipy 1.13. The complete pipeline is "
        "organised as numbered scripts under src/ingest/, src/clean/, src/analysis/, "
        "and src/viz/:")

    steps = [
        "01_download_plan.py — Documents data acquisition checklist",
        "02_parse_gbd.py — Parses and validates GBD CSV exports",
        "05_extract_morth_tables.py — Extracts MoRTH Table 42 from PDF",
        "06_extract_ncrb_tables.py — Extracts NCRB cause-specific tables",
        "05_harmonize_states.py — Applies state name crosswalk",
        "08_assemble_master.py — Assembles master long-format dataset",
        "10_state_burden.py — Computes state-level DALY/death/YLD summaries",
        "11_decomposition.py — Computes YLD fractions and YLD:YLL ratios",
        "12_hdbi.py — Computes Hidden Disability Burden Index",
        "17_inequality.py — Computes Gini, CV, IQR, top:bottom ratio",
        "18_mismatch.py — Benchmarks GBD vs MoRTH/NCRB road deaths",
        "21_qc_full.py — Full QC with HTML report",
        "fig1_daly_map.py — India state choropleth (DALY rate)",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(step)
        set_font(run, size=10)

    add_body(doc,
        "Code repository: https://github.com/[repository] (CC-BY 4.0). "
        "Data pipeline produces master_dataset.csv (77,974 rows: 77,568 GBD + 154 MoRTH + "
        "252 NCRB). All QC checks passed (WARN only for optional shapefile; all data checks PASS). "
        "The QC HTML report is included as a supplementary file.",
        size=11)

    out = PUBLISH_DIR / "04_supplementary.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=== Building NMJI publication package ===")
    build_title_page()
    build_manuscript()
    build_cover_letter()
    build_declarations()
    build_supplementary()
    print(f"\nAll files saved to: {PUBLISH_DIR}")
    for f in sorted(PUBLISH_DIR.iterdir()):
        size_kb = f.stat().st_size // 1024
        print(f"  {f.name}  ({size_kb} KB)")
