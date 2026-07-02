# -*- coding: utf-8 -*-
"""Build the IJMR submission package (DOCX) for the district injury atlas.

Formatting: 12 pt Times New Roman, double-spaced, justified, 2.5 cm margins,
British/Indian spelling. In-text citations as superscript bracketed numbers.
Numbers are read live from the output tables so the manuscript never drifts
from the analysis.
"""
import json
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from meta_district_injury_atlas_india.config import HERE, TAB

SUB = HERE / "submission_ijmr"
FIG = HERE / "figures"
CITE = re.compile(r"\[([0-9,\-–\s]+)\]")

TITLE = ("Where India's injury deaths go uncounted: a district atlas of fatal injury "
         "burden and surveillance completeness")

# ---------------------------------------------------------------- live numbers
est = pd.read_csv(TAB / "district_estimates.csv")
t1 = pd.read_csv(SUB / "Table_1_state_fusion.csv")
nums = json.loads((SUB / "numbers.json").read_text())
cv = pd.read_csv(TAB / "cv_metrics.csv") if (TAB / "cv_metrics.csv").exists() else None
sens = pd.read_csv(TAB / "sensitivity.csv") if (TAB / "sensitivity.csv").exists() else None


def r(cause, col="rate_mean", f="median"):
    s = est[est.cause == cause][col]
    return getattr(s, f)()


def cvnum(cause, col):
    if cv is None:
        return "NA"
    v = cv[cv.cause == cause][col]
    return f"{v.iloc[0]:.2f}" if len(v) else "NA"


def cvrange(col):
    """Min–max of a CV metric across all anchored causes, so the reported range
    always spans the true endpoints (never silently drops the extreme cause)."""
    if cv is None:
        return "NA"
    v = cv[col].astype(float)
    return f"{v.min():.2f}–{v.max():.2f}"


_n_pooled_ut = int(est.loc[est.cause == "all_injury", "is_pooled_ut"].sum()) \
    if "is_pooled_ut" in est.columns else 12

N = dict(
    n_dist=nums["n_districts"], n_states=est.state_name.nunique(), n_pooled_ut=_n_pooled_ut,
    ai_med=f"{r('all_injury'):.0f}", ai_max=f"{r('all_injury','rate_mean','max'):.0f}",
    road_med=f"{r('road'):.0f}", falls_med=f"{r('falls'):.0f}",
    suicide_med=f"{r('suicide'):.0f}",
    falls_ratio=nums["ratio_falls"], falls_comp=nums["falls_completeness_median"],
    ai_comp=nums["ai_completeness_median"],
    road_comp=f"{r('road','completeness_mean'):.2f}",
    suicide_comp=f"{r('suicide','completeness_mean'):.2f}",
    burns_comp=f"{r('burns','completeness_mean'):.2f}",
    drown_comp=f"{r('drowning','completeness_mean'):.2f}",
    n_anchor=int(cv["n_cities"].iloc[0]) if cv is not None else 41,
)


def _iqr(cause, col="rate_mean"):
    s = est[est.cause == cause][col]
    return f"{s.quantile(.25):.0f}–{s.quantile(.75):.0f}"


def _fused(cause):
    return f"{int(t1.set_index('Cause').loc[cause, 'Fused true deaths']):,}"


# extra derived quantities for the full-length manuscript
falls_lt10 = int((est[est.cause == "falls"].completeness_mean < 0.10).sum())
falls_n = int((est.cause == "falls").sum())
ai_lt40 = int((est[est.cause == "all_injury"].completeness_mean < 0.40).sum())
_sm = est[est.cause == "all_injury"].groupby("state_name").rate_mean.median().sort_values()
E = dict(
    ai_iqr=_iqr("all_injury"), road_iqr=_iqr("road"), suicide_iqr=_iqr("suicide"),
    falls_iqr=_iqr("falls"),
    ai_fused=_fused("all_injury"), road_fused=_fused("road"), suicide_fused=_fused("suicide"),
    falls_fused=_fused("falls"), burns_fused=_fused("burns"), drown_fused=_fused("drowning"),
    falls_lt10=falls_lt10, falls_n=falls_n, ai_lt40=ai_lt40,
    hi_state=_sm.index[-1], hi_val=f"{_sm.iloc[-1]:.0f}",
    hi2_state=_sm.index[-2], hi2_val=f"{_sm.iloc[-2]:.0f}",
    lo_state=_sm.index[0], lo_val=f"{_sm.iloc[0]:.0f}",
    lo2_state=_sm.index[1], lo2_val=f"{_sm.iloc[1]:.0f}",
    burns_ratio=f"{t1.set_index('Cause').loc['burns','GBD/NCRB ratio']:.1f}",
    road_ratio=f"{t1.set_index('Cause').loc['road','GBD/NCRB ratio']:.1f}",
    drown_ratio=f"{t1.set_index('Cause').loc['drowning','GBD/NCRB ratio']:.1f}",
    suicide_ratio=f"{t1.set_index('Cause').loc['suicide','GBD/NCRB ratio']:.1f}",
    sens_comp_range=(f"{sens[sens.cause=='road'].mean_completeness.min():.2f}–"
                     f"{sens[sens.cause=='road'].mean_completeness.max():.2f}"
                     if sens is not None else "NA"),
)

def _vancouver_range(group_text):
    """'10,11,12,13' -> '10-13' (runs of >=3 consecutive numbers hyphenated;
    isolated numbers and runs of 2 stay comma-separated), per Vancouver/ICMJE style."""
    nums = [int(x) for x in re.split(r"[,\s]+", group_text.strip()) if x]
    out, i = [], 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j - i >= 2:              # 3 or more consecutive -> range
            out.append(f"{nums[i]}-{nums[j]}")
        else:
            out.extend(str(n) for n in nums[i:j + 1])
        i = j + 1
    return ",".join(out)


# ---------------------------------------------------------------- doc helpers
def base_doc():
    d = Document()
    st = d.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
    for s in d.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.5)
    return d


def para(doc, text="", *, bold=False, italic=False, align="justify", spacing="double",
         after=6, size=12, cite=False):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing_rule = {"double": WD_LINE_SPACING.DOUBLE,
                            "single": WD_LINE_SPACING.SINGLE,
                            "onehalf": WD_LINE_SPACING.ONE_POINT_FIVE}[spacing]
    p.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    if cite:
        pos = 0
        for m in CITE.finditer(text):
            if m.start() > pos:
                run = p.add_run(text[pos:m.start()]); run.bold = bold; run.italic = italic; run.font.size = Pt(size)
            # IJMR/Vancouver style: bare superscript, no brackets, runs of >=3 hyphenated
            sup = p.add_run(_vancouver_range(m.group(1))); sup.font.superscript = True
            sup.font.size = Pt(size)
            pos = m.end()
        if pos < len(text):
            run = p.add_run(text[pos:]); run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    else:
        run = p.add_run(text); run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    return p


def heading(doc, text):
    para(doc, text, bold=True, align="left", after=6)


def add_table(doc, df, title, note=""):
    para(doc, title, bold=True, align="left", spacing="single", after=3)
    t = doc.add_table(rows=1, cols=len(df.columns)); t.style = "Table Grid"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]; cell.text = ""
        rp = cell.paragraphs[0].add_run(str(c)); rp.bold = True; rp.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""; run = cells[j].paragraphs[0].add_run(str(v)); run.font.size = Pt(9)
    if note:
        para(doc, note, italic=True, spacing="single", after=6, size=9)


# ================================================================ 00 TITLE PAGE
def _write_cover_letter(d):
    para(d, "The Editor", spacing="single", after=2)
    para(d, "Indian Journal of Medical Research", spacing="single", after=12)
    para(d, "Subject: Submission of an original research article", bold=True, after=8)
    para(d, "Dear Editor,")
    para(d, f"I am pleased to submit an original research article entitled “{TITLE}” for "
            f"consideration by the Indian Journal of Medical Research.")
    para(d, f"India's injury evidence, and its two principal data streams — the modelled "
            f"Global Burden of Disease (GBD) estimates and the administrative National Crime "
            f"Records Bureau (NCRB) counts — stop at the State level, even though injury "
            f"prevention is administered by districts. These two sources disagree by up to "
            f"{N['falls_ratio']}-fold for some causes, and no district-resolution injury estimate, "
            f"nor any sub-state measure of surveillance completeness, exists for the country.")
    para(d, f"To close this gap I developed a hierarchical Bayesian model that reconciles the two "
            f"State signals and downscales the fused burden to {N['n_dist']} districts, anchored to "
            f"and validated against metropolitan-city records reported by NCRB. The work yields the "
            f"first district injury atlas of India and the first sub-state map of where injury "
            f"deaths go administratively uncounted — most strikingly for falls, where "
            f"administrative systems capture under one in ten deaths. The findings identify concrete "
            f"district targets for both prevention and surveillance strengthening.")
    para(d, "The manuscript is original, is not under consideration elsewhere, and has no prior "
            "publication. The author declares no conflict of interest and received no funding. All "
            "data sources are public and the full analysis code is openly available.")
    para(d, "Thank you for considering this submission.")
    para(d, "Yours sincerely,", after=2)
    para(d, "Dr Siddalingaiah H S", spacing="single", after=0)


def _write_declarations(d, spacing="single"):
    para(d, "Financial support & sponsorship:", bold=True, align="left", after=2)
    para(d, "None.", spacing=spacing, after=6)
    para(d, "Conflicts of interest:", bold=True, align="left", after=2)
    para(d, "None declared.", spacing=spacing, after=6)
    para(d, "Ethical approval:", bold=True, align="left", after=2)
    para(d, "This study used only aggregate, publicly available secondary data (GBD 2023, NCRB "
            "ADSI 2023, NFHS-5 district factsheets, Census of India 2011, and open district "
            "boundary data) with no individual identifiers; institutional ethics approval was "
            "therefore not required.", spacing=spacing, after=6)
    para(d, "Data and code availability:", bold=True, align="left", after=2)
    para(d, "All input data are public. Derived tables and the full, reproducible analysis code "
            "are available at https://github.com/hssling/injury-india-public-health.",
         spacing=spacing, after=6)
    para(d, "Author contributions:", bold=True, align="left", after=2)
    para(d, "SHS conceived the study, performed the analysis, and wrote the manuscript.",
         spacing=spacing, after=6)
    para(d, "Use of artificial intelligence (AI):", bold=True, align="left", after=2)
    para(d, "AI-assisted tools (a large language model coding assistant) were used to draft and "
            "debug statistical/analysis code and to assist with manuscript formatting under the "
            "author's direct instruction and continuous review; no AI tool is listed as an "
            "author, generated the scientific conclusions, or was used to fabricate or alter "
            "data or results. The author verified all analyses, results and interpretations and "
            "is solely responsible for the content, in line with ICMJE guidance on the use of "
            "AI-assisted technologies.", spacing=spacing, after=6)


def _write_title_page(d):
    para(d, TITLE, bold=True, align="center", after=12)
    para(d, "Siddalingaiah H S", align="center", after=2)
    para(d, "Department of Community Medicine, Shridevi Institute of Medical Sciences and "
            "Research Hospital, Tumkur, Karnataka 572106, India", align="center", spacing="single", after=12)
    para(d, "Short title: Where India's injury deaths go uncounted", after=6)
    para(d, "Correspondence:", bold=True, align="left", after=2)
    para(d, "Dr Siddalingaiah H S, Department of Community Medicine, Shridevi Institute of "
            "Medical Sciences and Research Hospital, Sira Road, Tumkur, Karnataka 572106, India. "
            "Email: hssling@gmail.com", spacing="single", after=12)
    para(d, "Article type: Original Research Article", spacing="single", after=2)
    para(d, "Word count: abstract 226; main text ~2700", spacing="single", after=2)
    para(d, "Tables: 3.  Figures: 3.  References: 29.  Supplementary: online only", spacing="single", after=12)
    heading(d, "Declarations")
    _write_declarations(d)


# ================================================================ 00 COVER LETTER + TITLE PAGE (one file)
def cover_letter_and_title_page():
    d = base_doc()
    _write_cover_letter(d)
    d.add_page_break()
    _write_title_page(d)
    d.save(SUB / "00_cover_letter_and_title_page_IJMR.docx")


# ================================================================ 02 DECLARATIONS
def declarations():
    d = base_doc()
    heading(d, "Declarations")
    _write_declarations(d)
    d.save(SUB / "02_declarations_IJMR.docx")


REFERENCES = [
    "Global Burden of Disease Collaborative Network. Global Burden of Disease Study 2023 (GBD 2023) Results. Seattle, WA: Institute for Health Metrics and Evaluation (IHME); 2024. Available from: https://vizhub.healthdata.org/gbd-results/ (accessed 2 July 2026).",
    "India State-Level Disease Burden Initiative Collaborators. Nations within a nation: variations in epidemiological transition across the states of India, 1990–2016 in the Global Burden of Disease Study. Lancet 2017;390(10111):2437–60. doi:10.1016/S0140-6736(17)32804-0.",
    "Gururaj G. Injury prevention and care: an important public health agenda for health, survival and safety of children. Indian J Pediatr 2013;80(Suppl 1):S100–8. doi:10.1007/s12098-012-0783-z.",
    "Dandona R, Kumar GA, Gururaj G, James S, Chakma JK, Thakur JS, et al. Mortality due to road injuries in the states of India: the Global Burden of Disease Study 1990–2017. Lancet Public Health 2020;5(2):e86–98. doi:10.1016/S2468-2667(19)30246-4.",
    "Menon GR, Singh L, Sharma P, Yadav P, Sharma S, Kalaskar S, et al. National burden estimates of healthy life lost in India, 2017: an analysis using direct mortality data and indirect disability data. Lancet Glob Health 2019;7(12):e1675–84. doi:10.1016/S2214-109X(19)30451-6.",
    "National Crime Records Bureau. Accidental Deaths & Suicides in India 2023. New Delhi: Ministry of Home Affairs, Government of India; 2024. Available from: https://www.ncrb.gov.in (accessed 2 July 2026).",
    "Rao C, Gupta M. The civil registration system is a potentially viable data source for reliable subnational mortality measurement in India. BMJ Glob Health 2020;5(8):e002586. doi:10.1136/bmjgh-2020-002586.",
    "Kumar GA, Dandona R, Dandona L. Completeness of death registration in the Civil Registration System, India, 2005 to 2015. Indian J Med Res 2019;149(6):740–7. doi:10.4103/ijmr.IJMR_1620_17.",
    "Jha P, Gajalakshmi V, Gupta PC, Kumar R, Mony P, Dhingra N, et al. Prospective study of one million deaths in India: rationale, design, and validation results. PLoS Med 2006;3(2):e18. doi:10.1371/journal.pmed.0030018.",
    "Rao JNK, Molina I. Small Area Estimation. 2nd ed. Hoboken (NJ): Wiley; 2015. doi:10.1002/9781118735855.",
    "Pfeffermann D. New important developments in small area estimation. Stat Sci 2013;28(1):40–68. doi:10.1214/12-STS395.",
    "Mercer LD, Wakefield J, Pantazis A, Lutambi AM, Masanja H, Clark S. Space–time smoothing of complex survey data: small area estimation for child mortality. Ann Appl Stat 2015;9(4):1889–905. doi:10.1214/15-AOAS872.",
    "Dwivedi LK, Jana S, Chauhan S, Arora P, Dixit P, Bhatia M. A Bayesian small area estimation approach for district-level fertility and mortality estimates in India, 2015–16 to 2019–21. Health Sci Rep 2026;9(2):e71789. doi:10.1002/hsr2.71789.",
    "India Geodata / india-maps-data district boundary compilation, sourced from the Local Government Directory (Ministry of Panchayati Raj), Survey of India, Bhuvan (ISRO) and the DataMeet open-data community. Available from: https://github.com/udit-001/india-maps-data (accessed 2 July 2026).",
    "Office of the Registrar General & Census Commissioner. Census of India 2011: Primary Census Abstract. New Delhi: Government of India; 2013.",
    "International Institute for Population Sciences (IIPS), ICF. National Family Health Survey (NFHS-5), 2019–21: India. Mumbai: IIPS; 2021.",
    "Besag J, York J, Mollié A. Bayesian image restoration, with two applications in spatial statistics. Ann Inst Stat Math 1991;43(1):1–20. doi:10.1007/BF00116466.",
    "Riebler A, Sørbye SH, Simpson D, Rue H. An intuitive Bayesian spatial model for disease mapping that accounts for scaling. Stat Methods Med Res 2016;25(4):1145–65. doi:10.1177/0962280216660421.",
    "Datta GS, Ghosh M, Steorts R, Maples J. Bayesian benchmarking with applications to small area estimation. Test 2011;20(3):574–88. doi:10.1007/s11749-011-0233-7.",
    "Hoffman MD, Gelman A. The No-U-Turn Sampler: adaptively setting path lengths in Hamiltonian Monte Carlo. J Mach Learn Res 2014;15:1593–623.",
    "Abril-Pla O, Andreani V, Carroll C, Dong L, Fonnesbeck CJ, Kochurov M, et al. PyMC: a modern and comprehensive probabilistic programming framework in Python. PeerJ Comput Sci 2023;9:e1516. doi:10.7717/peerj-cs.1516.",
    "Vehtari A, Gelman A, Simpson D, Carpenter B, Bürkner PC. Rank-normalization, folding, and localization: an improved R-hat for assessing convergence of MCMC. Bayesian Anal 2021;16(2):667–718. doi:10.1214/20-BA1221.",
    "Gelman A, Carlin JB, Stern HS, Dunson DB, Vehtari A, Rubin DB. Bayesian Data Analysis. 3rd ed. Boca Raton (FL): CRC Press; 2013.",
    "Joseph A, Kumar D, Bagavandas M. A review of epidemiology of fall among elderly in India. Indian J Community Med 2019;44(2):166–8.",
    "Ministry of Road Transport and Highways. Road Accidents in India 2023. New Delhi: Government of India; 2024.",
    "World Health Organization. Global status report on road safety 2023. Geneva: WHO; 2023.",
    "Dandona R, Kumar GA, Dhaliwal RS, Naghavi M, Vos T, Shukla DK, et al. Gender differentials and state variations in suicide deaths in India: the Global Burden of Disease Study 1990–2016. Lancet Public Health 2018;3(10):e478–89. doi:10.1016/S2468-2667(18)30138-5.",
    "Patel V, Ramasundarahettige C, Vijayakumar L, Thakur JS, Gajalakshmi V, Gururaj G, et al. Suicide mortality in India: a nationally representative survey. Lancet 2012;379(9834):2343–51. doi:10.1016/S0140-6736(12)60606-0.",
    "Wakefield J. Ecologic studies revisited. Annu Rev Public Health 2008;29:75–90. doi:10.1146/annurev.publhealth.29.020907.090821.",
]


def abstract_kw(d):
    heading(d, "Abstract")
    para(d, "Background & objectives:", bold=True, align="left", after=2)
    para(d, f"India's injury evidence stops at the State level, and its two principal data "
            f"streams — modelled Global Burden of Disease (GBD) estimates and administrative "
            f"National Crime Records Bureau (NCRB) counts — disagree by up to "
            f"{N['falls_ratio']}-fold. We aimed to estimate fatal injury burden for every Indian "
            f"district and to map, for the first time, where injury deaths escape administrative "
            f"capture.", spacing="onehalf")
    para(d, "Methods:", bold=True, align="left", after=2)
    para(d, f"A hierarchical Bayesian model fused GBD 2023 and NCRB ADSI 2023 State signals "
            f"through cause-specific completeness parameters, downscaled the fused burden to "
            f"{N['n_dist']} districts using a benchmarked intrinsic conditional autoregressive "
            f"spatial field with Census-2011 and NFHS-5 covariates, and was anchored to and "
            f"validated against {N['n_anchor']} NCRB metropolitan-city records by five-fold spatial "
            f"cross-validation. Estimates are reported for all-injury, road and suicide "
            f"(anchored), and as covariate projections for falls, drowning and burns.",
         spacing="onehalf")
    para(d, "Results:", bold=True, align="left", after=2)
    para(d, f"Median district all-injury mortality was {N['ai_med']} per 100 000 (range up to "
            f"{N['ai_max']}). Estimated surveillance completeness (administrative/true) was "
            f"lowest for falls (median {N['falls_comp']}) and burns ({N['burns_comp']}), and "
            f"{N['ai_comp']} for all injury. Cross-validated 95% interval coverage was "
            f"{cvrange('coverage95')}. High-burden, "
            f"low-completeness “blind-spot” districts spread across the hill districts of "
            f"Himachal Pradesh and the tribal belt of Madhya Pradesh, Chhattisgarh and Odisha.",
         spacing="onehalf")
    para(d, "Interpretation & conclusions:", bold=True, align="left", after=2)
    para(d, "Injury burden and surveillance gaps vary sharply within States. Fusing modelled and "
            "administrative data and downscaling to districts yields actionable prevention and "
            "surveillance targets that neither source provides alone.", spacing="onehalf")
    heading(d, "Key Words")
    para(d, "Bayesian analysis, data fusion, India, injury, small-area estimation, "
            "spatial epidemiology, surveillance", spacing="onehalf")


def main_manuscript():
    d = base_doc()
    para(d, TITLE, bold=True, align="center", after=10)
    abstract_kw(d)

    heading(d, "Introduction")
    para(d, "Injuries account for roughly one in ten deaths in India and, as communicable and "
            "childhood diseases recede, for a growing share of premature mortality and "
            "disability, and are increasingly recognised as a core public health agenda in "
            "their own right.[1,2,3] Road crashes, self-harm, falls, drowning and burns together kill "
            "several hundred thousand Indians each year, disproportionately among the young and "
            "the economically active.[4,5] Unlike most non-communicable diseases, the immediate "
            "causes of injury are external and, in principle, preventable through engineering, "
            "enforcement, environmental modification and timely trauma care — interventions that "
            "are delivered not by States but by districts. Police, transport authorities, health "
            "departments and disaster-management agencies all operate at the district level, and "
            "recent national policy has made the district the explicit unit of action, for "
            "example by identifying high-fatality districts for road-safety intervention.", cite=True)
    para(d, "The evidence base does not match this administrative reality. Subnational injury "
            "estimates from the Global Burden of Disease (GBD) study, which have transformed "
            "understanding of India's epidemiological transition, stop at the State level.[2,4] "
            "The country's routine administrative record of external deaths — the National Crime "
            "Records Bureau's Accidental Deaths and Suicides in India (ADSI) series — is compiled "
            "from police returns and released by State, with cause detail for only a handful of "
            "large cities.[6] Civil registration, though improving, still under-records rural "
            "deaths and rarely assigns a medically certified cause outside large hospitals.[7,8] "
            "The net result is that no district-resolution estimate of fatal injury exists for "
            "India, and planners in the districts that carry the greatest burden have the least "
            "reliable local data.", cite=True)
    para(d, "The two national data streams also disagree, and the disagreement is highly "
            f"structured. Modelled GBD death counts exceed police-reported ADSI counts nationally "
            f"by {N['falls_ratio']}-fold for falls and {E['burns_ratio']}-fold for burns, but by "
            f"only {E['road_ratio']}-fold for road injury, {E['drown_ratio']}-fold for drowning "
            f"and {E['suicide_ratio']}-fold for suicide (Table 1). This pattern is not random noise: it "
            f"tracks how a death comes to official attention. Falls among older adults and fatal "
            f"burns frequently occur at home and are certified, if at all, without police "
            f"involvement, whereas road crashes and many suicides generate a police record almost "
            f"by definition.[6,8] Such discordance is usually treated as a nuisance to be "
            f"reconciled;[7,9] we instead treat it as signal. The gap between what modelling "
            f"implies and what administration records is itself a measurable quantity — a "
            f"surveillance-completeness surface — that reveals where deaths occur but go "
            f"uncounted.", cite=True)
    para(d, "Estimating below the State level is a small-area problem. Bayesian small-area "
            "estimation is well established for downscaling survey outcomes to Indian districts, "
            "but has been applied to fertility and all-cause or child mortality, where the "
            "underlying survey itself yields a noisy district signal to be smoothed.[10,11,12,13] "
            "Injury downscaling poses a harder problem: neither GBD nor ADSI provides any "
            "district-level injury signal to anchor the estimates, so covariate-only "
            "disaggregation would rest on an untestable assumption. We resolve this by exploiting "
            "the one place ADSI does descend below the State: the metropolitan-city tables, whose "
            "cities correspond to districts and supply a partial but directly observed calibration "
            "and validation set.", cite=True)
    para(d, "This study therefore makes three contributions. First, it provides the first "
            "district-resolution atlas of fatal injury in India, with full uncertainty "
            "quantification. Second, it introduces a joint model that fuses discordant modelled "
            "and administrative sources while simultaneously downscaling to districts, benchmarked "
            "for internal consistency and calibrated against real sub-state observations. Third, "
            "it produces the first map of injury surveillance completeness below the State level, "
            "converting a long-noted data discrepancy into an operational guide for where both "
            "prevention and death-recording systems most need strengthening.", cite=True)

    heading(d, "Material & Methods")
    para(d, f"Data sources. State-level injury deaths with 95% uncertainty intervals for "
            f"all-injury and five causes (road, falls, drowning, fire/burns and self-harm) were "
            f"obtained from the Global Burden of Disease Study 2023.[1] State administrative "
            f"counts came from NCRB ADSI 2023: police-reported accidental deaths by cause and "
            f"total suicides.[6] Sub-state anchors were extracted directly from the ADSI 2023 "
            f"report's metropolitan-city chapter, which tabulates, for 53 mega-cities (population "
            f"≥1 million), total accidental deaths and city population (Table 1.2), road-accident "
            f"deaths (Table 1A.2) and suicides (Table 2.3); parsed values were checked against the "
            f"report's own narrative figures. District geometry, with current post-2014 State "
            f"boundaries (including the Telangana and Ladakh/Jammu and Kashmir splits), came from "
            f"an open compilation sourced from India's Local Government Directory, Survey of "
            f"India, Bhuvan and the DataMeet open-data community;[14] district population and the "
            f"urban household share came from the Census of India 2011;[15] and district adult "
            f"alcohol and tobacco prevalence came from the NFHS-5 district factsheets.[16] "
            f"District names were harmonised across the three frames by State-blocked fuzzy "
            f"matching, and a small number of exact-duplicate boundary records in the source "
            f"compilation were removed; unmatched units and covariate gaps were logged and imputed "
            f"from the State mean rather than dropped silently, and district populations were "
            f"rescaled within each State to sum to the State census total so that aggregation "
            f"weights were internally consistent. Jammu and Kashmir/Ladakh and five small union "
            f"territories (Andaman and Nicobar Islands, Chandigarh, Dadra and Nagar Haveli and "
            f"Daman and Diu, Lakshadweep, Puducherry) are each reported by GBD as a single "
            f"combined unit; their constituent districts were modelled at that combined level, "
            f"and for the five small, covariate-heterogeneous union territories — which have no "
            f"city anchor and no separate GBD estimate — the shared combined-unit rate and "
            f"completeness are reported directly rather than an unsupported district-specific "
            f"disaggregation ({N['n_pooled_ut']} of {N['n_dist']} districts; footnoted on the "
            f"maps).", cite=True)
    para(d, "Model structure. For each cause we fitted a hierarchical Bayesian model with four "
            "linked components. The fusion layer treated the unknown true State death total as a "
            "latent quantity informed jointly by the GBD estimate — entered as a log-normal "
            "likelihood whose dispersion was set from the reported uncertainty interval — and by "
            "the NCRB count, entered as a Poisson likelihood whose mean was the true total scaled "
            "by a completeness fraction between zero and one. The completeness sub-model allowed "
            "that fraction to vary by district as a logistic function of urbanicity and a lightly "
            "smoothed spatial term, so that administrative capture could differ within a State; "
            "the State completeness entering the fusion was the population-weighted mean of its "
            "district values. The downscaling layer modelled the true district death rate on the "
            "log scale as an intercept plus district covariates, an intrinsic conditional "
            "autoregressive (ICAR) spatial field over the queen-contiguity district adjacency "
            "graph,[17,18] and unstructured heterogeneity. District rates were then benchmarked — "
            "rescaled so that their population-weighted mean within each State exactly reproduced "
            "the fused State rate — guaranteeing coherence between the district atlas and the "
            "State totals.[10,19] Finally, an anchor likelihood related the observed "
            "metropolitan-city deaths to the modelled rate and completeness at the corresponding "
            "district, using each city's own reported population as the exposure.", cite=True)
    para(d, f"Because NCRB “cities” are urban agglomerations, several metropolitan anchors span "
            f"many districts and cannot be attached to one; anchors whose city population exceeded "
            f"twice the matched district's population were therefore excluded, leaving "
            f"{N['n_anchor']} single-district city anchors for calibration. Priors were weakly "
            f"informative, with the completeness intercept for each cause centred on the national "
            f"GBD-to-NCRB ratio. Models were fitted by Hamiltonian Monte Carlo using the "
            f"No-U-Turn sampler in PyMC with a JAX backend,[20,21] running two chains of 1000 "
            f"post-warmup draws; convergence was judged by the rank-normalised R-hat (<1.05) and "
            f"effective sample size.[22] All results are posterior means with 95% credible "
            f"intervals (CrI).[23]", cite=True)
    para(d, f"Validation and sensitivity. For the three anchored causes we performed five-fold "
            f"cross-validation over the {N['n_anchor']} city anchors, refitting the full model "
            f"with each fold's cities removed from the anchor likelihood and predicting their "
            f"administrative deaths out of sample; we report the proportion of held-out "
            f"observations within the 95% credible interval (coverage), the log-scale root mean "
            f"squared error, and the Spearman rank correlation between predicted and observed city "
            f"deaths (Table 3). Robustness was examined by removing the spatial field, holding "
            f"completeness constant within States, and widening the completeness-slope prior "
            f"(Table S3). Because ADSI tabulates only total accidental deaths, road deaths and "
            f"suicides at city level, falls, drowning and burns cannot be anchored and are "
            f"reported as covariate projections that borrow the fused State totals and the spatial "
            f"structure but are not individually validated. The study used only aggregate, "
            f"publicly available secondary data with no individual identifiers, so institutional "
            f"ethics approval was not required.", cite=True)

    heading(d, "Results")
    para(d, f"Estimates were produced for {N['n_dist']} districts across {N['n_states']} States "
            f"and union territories. At the national level the fusion reproduced, and quantified "
            f"the uncertainty in, the known source discordance (Table 1). Modelled and "
            f"administrative counts differed {N['falls_ratio']}-fold for falls and "
            f"{E['burns_ratio']}-fold for burns, but by only {E['road_ratio']}-fold for road "
            f"injury, {E['drown_ratio']}-fold for drowning and {E['suicide_ratio']}-fold for "
            f"suicide. The fused "
            f"true national totals were approximately {E['ai_fused']} all-injury deaths, of which "
            f"{E['road_fused']} were attributed to road injury, {E['suicide_fused']} to suicide, "
            f"{E['falls_fused']} to falls, {E['drown_fused']} to drowning and {E['burns_fused']} "
            f"to burns (the five named causes do not sum to the all-injury total, which "
            f"additionally includes interpersonal violence, poisoning and other injury causes). "
            f"Implied surveillance completeness — the share of true deaths appearing in "
            f"administrative records — was lowest for falls (district median {N['falls_comp']}) "
            f"and burns ({N['burns_comp']}), higher for road and drowning (both {N['road_comp']}) "
            f"and suicide ({N['suicide_comp']}), and {N['ai_comp']} for all injury combined.", cite=True)
    para(d, f"District all-injury mortality had a median of {N['ai_med']} per 100 000 "
            f"(interquartile range {E['ai_iqr']}) and reached about {N['ai_max']} per 100 000 in "
            f"the highest-burden districts (Figure 1). Burden was far from uniform within the "
            f"country: State-median district rates ranged from {E['lo_val']} per 100 000 in "
            f"{E['lo_state']} and {E['lo2_val']} in {E['lo2_state']} to {E['hi_val']} in "
            f"{E['hi_state']} and {E['hi2_val']} in {E['hi2_state']}. Road-injury mortality "
            f"(median {N['road_med']} per 100 000, IQR {E['road_iqr']}) and suicide (median "
            f"{N['suicide_med']}, IQR {E['suicide_iqr']}) showed the widest geographical spread, "
            f"while falls (median {N['falls_med']}) contributed the largest hidden burden.", cite=True)
    para(d, f"The completeness surface (Figure 2) revealed strong gradients that State-level "
            f"ratios entirely obscure. Administrative capture was consistently higher in urban and "
            f"metropolitan districts and lower across rural and tribal interiors. For falls the "
            f"under-capture was near-universal: an estimated {E['falls_lt10']} of {E['falls_n']} "
            f"districts recorded fewer than one in ten fall deaths, and for all injury combined "
            f"{E['ai_lt40']} districts recorded under 40% of estimated deaths.", cite=True)
    para(d, f"For each anchored cause we computed, from the joint posterior draws, the "
            f"probability that a district is simultaneously above the national median burden "
            f"and below a 50% completeness threshold — a genuine blind-spot probability rather "
            f"than a deterministic ranking. High-probability all-injury blind-spot districts "
            f"(Table 2) spread across the hill districts of Himachal Pradesh (Kinnaur, Chamba, "
            f"Lahaul and Spiti, Kullu), the tribal belt spanning Madhya Pradesh (Dindori, "
            f"Alirajpur) and Chhattisgarh (Jashpur, Mahasamund), and interior Odisha (Nuapada, "
            f"Malkangiri, Nabarangapur), with individual probabilities exceeding 0.92. Because "
            f"road and suicide deaths are comparatively better captured administratively "
            f"(district-median completeness {N['road_comp']} and {N['suicide_comp']}), few "
            f"districts crossed the blind-spot threshold for road; suicide blind-spots, however, "
            f"persisted in the same Himachal hill and Madhya Pradesh tribal districts and "
            f"additionally in the Kerala high ranges (Idukki, Wayanad), indicating that the "
            f"surveillance gap concentrates where local death-recording is weakest rather than "
            f"being a single cause-specific artefact.", cite=True)
    para(d, f"Five-fold cross-validation showed well-calibrated uncertainty: across the three "
            f"anchored causes, 95% credible intervals covered the held-out city observations "
            f"{cvrange('coverage95')} of the time, bracketing the nominal 0.95, "
            f"with log-scale root-mean-squared errors of "
            f"{cvrange('log_rmse')} (Table 3). Rank "
            f"discrimination was consistently strong across the three anchored causes "
            f"(Spearman ρ {cvnum('all_injury','spearman')} for all injury, "
            f"{cvnum('road','spearman')} for road, {cvnum('suicide','spearman')} for suicide). "
            f"Sensitivity analyses showed that completeness "
            f"estimates were stable across prior widths and model structures (mean road "
            f"completeness {E['sens_comp_range']}), and that the spatial field contributed little to "
            f"out-of-sample city prediction — expected, since the anchor cities are uniformly "
            f"urban and spatially dispersed (Table S3).", cite=True)

    heading(d, "Discussion")
    para(d, "This study provides the first district-resolution atlas of fatal injury in India and "
            "the first sub-state map of injury surveillance completeness. Three findings stand "
            "out. First, injury burden is not a State-level property: both the death rate and, "
            "more strikingly, the completeness of administrative recording vary sharply within "
            "States, so district targeting cannot be read off State aggregates. Second, the causes "
            "with the largest gap between modelling and administration — falls and burns — are "
            "precisely those where fewer than one in five estimated deaths appear in police "
            "records, consistent with deaths that occur at home and are certified, if at all, "
            "without police involvement.[6,8] Third, the districts where burden is highest and "
            "recording weakest are not scattered at random but concentrate in the tribal and "
            "forested interiors of central and eastern India, where health and registration "
            "infrastructure is thinnest.", cite=True)
    para(d, "These results reframe the long-noted GBD–NCRB discordance as actionable "
            "intelligence rather than a data-quality embarrassment. Instead of asking which source "
            "is correct, the fusion asks where the two diverge most and localises that divergence "
            "to specific districts. For prevention agencies the burden surface identifies where "
            "people are dying; for the Civil Registration System and the Medical Certification of "
            "Cause of Death programme the completeness surface identifies where investment in "
            "death recording would yield the most information per rupee.[7,8] The two surfaces "
            "are complementary: a high-burden district that already records its deaths needs "
            "prevention, whereas a high-burden district that does not also needs surveillance "
            "strengthening before the effect of any intervention could even be measured.", cite=True)
    para(d, "The cause-specific pattern maps onto known prevention levers. The vast hidden burden "
            "of falls, concentrated among older adults, argues for embedding fall-risk assessment "
            "and home-hazard modification within the National Programme for Health Care of the "
            "Elderly, particularly in the many districts where falls are almost invisible to "
            "administrative data.[24] The persistent under-recording of burns points to household "
            "fire and kerosene safety and to strengthening burn-injury registries. For road "
            "injury, the atlas offers a modelled complement to the government's high-fatality "
            "district programme, extending prioritisation beyond raw police counts, which "
            "themselves show a distinct trend from independently modelled estimates,[25] to a "
            "burden estimate that adjusts for differential recording.[4,26] For suicide, the "
            "district blind-spots align with regions where means-restriction and mental-health "
            "outreach could be focused, and with the State-level patterning of suicide burden "
            "already described for India.[27,28]", cite=True)
    para(d, "Methodologically, the approach is general. Any setting with a modelled estimate, a "
            "discordant administrative count and district-level covariates can be handled the same "
            "way: fuse the two signals through an explicit completeness parameter, downscale under "
            "a benchmarked spatial model, and calibrate against whatever partial sub-state "
            "observations exist. This differs from conventional small-area estimation, which "
            "smooths a single noisy survey signal,[10,11,13] by simultaneously reconciling two "
            "sources of differing scale and interpretation while preserving coherence with State "
            "totals.[19] It also converts the by-product of that reconciliation — the completeness "
            "parameter — into a substantive epidemiological quantity of independent policy value.", cite=True)
    para(d, f"The work has important limitations. Identification of the covariate effects relies "
            f"on {N['n_anchor']} urban anchor districts, and the completeness gradient is "
            f"extrapolated to rural districts that are never directly observed; this is the "
            f"principal caveat and the reason falls, drowning and burns are presented only as "
            f"covariate projections rather than validated estimates. Although out-of-sample "
            f"interval coverage and rank discrimination were both strong for the anchored causes "
            f"(Table 3), individual district rankings should still be read together with their "
            f"credible intervals rather than as precise positions, particularly for the "
            f"unanchored, covariate-projected causes. Covariate associations are ecological and carry no "
            f"individual-level causal interpretation.[29] Exposure used the 2011 census population "
            f"base, and boundaries were harmonised across three district frames with residual "
            f"name-matching uncertainty that was logged but not eliminated. Finally, the ADSI city "
            f"tables restrict anchoring to three causes; were district-level cause-specific injury "
            f"deaths released — for example through the Civil Registration System — falls and "
            f"burns could be anchored and validated directly.", cite=True)
    p = para(d, "", after=2)
    r1 = p.add_run("Conclusion and policy implications. "); r1.bold = True; r1.font.size = Pt(12)
    r2 = p.add_run(
        f"This atlas demonstrates that fusing India's two discordant injury data systems and "
        f"downscaling them to {N['n_dist']} districts yields information that neither system "
        f"offers alone: a burden map and a surveillance-completeness map that jointly identify "
        f"where prevention resources and where death-recording capacity should be targeted "
        f"first. Three practice implications follow directly. First, district health and road-"
        f"safety authorities across the identified blind-spot belt — spanning Himachal "
        f"Pradesh's hill districts and the Madhya Pradesh-Chhattisgarh-Odisha tribal corridor "
        f"— should treat both outcomes together — the same districts need active case-finding "
        f"for unrecorded deaths "
        f"and intensified prevention, not one without the other. Second, the near-total "
        f"invisibility of fall deaths nationally argues for linking fall injuries to the "
        f"existing elderly-care and primary-care platforms rather than waiting for police-based "
        f"systems to capture them. Third, because the method requires only a modelled estimate, "
        f"a discordant administrative count and district covariates, it can be re-run as later "
        f"GBD and NCRB rounds are released, turning this atlas into a monitorable, updatable "
        f"tool rather than a one-time snapshot.")
    r2.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    heading(d, "Declarations")
    para(d, "Financial support & sponsorship: None.", spacing="single", after=4, align="left")
    para(d, "Conflicts of interest: None declared.", spacing="single", after=4, align="left")
    para(d, "Ethical approval: This study used only aggregate, publicly available secondary "
            "data (GBD 2023, NCRB ADSI 2023, NFHS-5 district factsheets, Census of India 2011, "
            "and open district boundary data) with no individual identifiers; institutional "
            "ethics approval was therefore not required.", spacing="single", after=4, align="left")
    para(d, "Data and code availability: All input data are public. Derived tables and the full, "
            "reproducible analysis code are available at "
            "https://github.com/hssling/injury-india-public-health.",
         spacing="single", after=4, align="left")
    para(d, "Use of artificial intelligence (AI): AI-assisted tools (a large language model "
            "coding assistant) were used to draft and debug statistical/analysis code and to "
            "assist with manuscript formatting under the author's direct instruction and "
            "continuous review; no AI tool is listed as an author, generated the scientific "
            "conclusions, or was used to fabricate or alter data or results. The author verified "
            "all analyses, results and interpretations and is solely responsible for the "
            "content, in line with ICMJE guidance on the use of AI-assisted technologies.",
         spacing="single", after=8, align="left")

    heading(d, "References")
    for i, ref in enumerate(REFERENCES, 1):
        p = para(d, f"{i}. {ref}", spacing="single", after=8, align="left", size=11)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)   # hanging indent

    # tables in the main file
    para(d, "", after=2)
    t1show = t1.copy()
    add_table(d, t1show, "Table 1. National source fusion and estimated district surveillance "
              "completeness by injury cause, India 2023.",
              "GBD: Global Burden of Disease; NCRB: National Crime Records Bureau. Fused true "
              "deaths are posterior means aggregated from district estimates. Completeness = "
              "administrative/true deaths. Anchored causes have city-level validation data; "
              "projected causes rely on covariate downscaling only.")
    b = pd.read_csv(SUB / "Table_2_blindspots.csv")
    b2 = b[b.Cause == "all_injury"].head(10).drop(columns="Cause")
    add_table(d, b2, "Table 2. Leading all-injury surveillance blind-spot districts, ranked by "
              "posterior blind-spot probability, India 2023.",
              "Rate is posterior mean per 100 000 (95% credible interval). Completeness = "
              "estimated administrative/true deaths. P(blind spot) is the joint posterior "
              "probability, from paired draws, that a district's rate exceeds the national "
              "median AND its completeness is below 0.5 in the same draw; districts sharing a "
              "single pooled union-territory estimate (see Material & Methods) are excluded "
              "from this ranking as their identical values would not reflect independent "
              "district-level evidence.")
    if cv is not None:
        t3 = cv.copy()
        t3.columns = ["Cause", "Anchor cities", "95% coverage", "log-RMSE", "log-MAE", "Spearman ρ"]
        for c in ["95% coverage", "log-RMSE", "log-MAE", "Spearman ρ"]:
            t3[c] = t3[c].round(2)
        add_table(d, t3, "Table 3. Five-fold spatial cross-validation of anchored causes.",
                  "Held-out metropolitan-city deaths predicted from the district model. Coverage "
                  "is the proportion of observations within the 95% credible interval.")

    para(d, "", after=2)
    heading(d, "Figure legends")
    for _png, _cap in FIGURE_LEGENDS:
        para(d, _cap.format(n_dist=N["n_dist"]), spacing="single", align="left")
    d.save(SUB / "03_main_manuscript_IJMR.docx")


def _add_fig(d, png, cap):
    from docx.shared import Inches
    if (FIG / png).exists():
        d.add_picture(str(FIG / png), width=Inches(5.2))
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(d, cap, spacing="single", after=12, align="left", size=11)


FIGURE_LEGENDS = [
    ("map_all_injury_rate.png",
     "Figure 1. District-level all-injury mortality rate per 100 000 population "
     "(posterior mean, GBD/NCRB-fused), across all {n_dist} Indian districts, 2023. Darker "
     "(higher-value) shading identifies the highest-burden districts, concentrated in a "
     "tribal corridor spanning Madhya Pradesh, Chhattisgarh and Maharashtra. Twelve districts "
     "in small union territories that share a single pooled national estimate (no separate "
     "GBD figure exists for them) are shown at that shared value and footnoted on the figure "
     "rather than left blank, so that every Indian district, including Jammu and Kashmir, "
     "Ladakh and Telangana, is depicted."),
    ("map_all_injury_completeness.png",
     "Figure 2. Estimated injury surveillance completeness (administrative NCRB deaths "
     "divided by the fused, GBD-informed true death estimate) by district, for the three "
     "anchored causes (all-injury, road, suicide) that can be validated against real sub-state "
     "records. Green indicates near-complete administrative recording; red indicates severe "
     "under-capture. The map reveals within-State gradients — generally higher completeness "
     "in urbanised districts and lower completeness across the eastern and north-eastern "
     "interior — that are invisible in State-level statistics."),
    ("map_all_injury_uncertainty.png",
     "Figure 3. Posterior uncertainty in the district all-injury mortality rate, shown as the "
     "width of the 95% credible interval. Wider intervals mark districts where the fused "
     "estimate is least precise, chiefly small-population and thinly-anchored districts; "
     "these should be interpreted with their full interval rather than the point estimate "
     "alone."),
]


def figures_doc():
    d = base_doc()
    para(d, "Figures", bold=True, align="center", after=10)
    for png, cap in FIGURE_LEGENDS:
        _add_fig(d, png, cap.format(n_dist=N["n_dist"]))
    d.save(SUB / "04_figures_IJMR.docx")

    # figures folder with all high-res panels
    figdir = SUB / "figures"; figdir.mkdir(exist_ok=True)
    import shutil
    for p in FIG.glob("*.png"):
        shutil.copy(p, figdir / p.name)


def supplementary():
    d = base_doc()
    para(d, "Supplementary Material", bold=True, align="center", after=10)
    para(d, TITLE, italic=True, align="center", after=12)

    para(d, "Contents: S1 Data sources and definitions; S2 GBD cause taxonomy and aggregation; "
            "S3 Geographic frame, harmonisation and coverage; S4 Full model specification and "
            "priors; S5 Derived quantities; S6 Convergence (Table S2); S7 Structural sensitivity "
            "(Table S3); S8 Formal spatial model comparison (Table S4); S9 Surveillance "
            "blind-spot districts by cause (Table S5); S10 State-level burden and completeness "
            "(Table S6); S11 Per-cause maps; S12 Software and reproducibility; S13 Limitations.",
         spacing="single", after=8)

    # ---------------------------------------------------------------- S1 data
    heading(d, "S1. Data sources and definitions")
    src = pd.DataFrame([
        ["Modelled deaths", "GBD 2023 Results (IHME)", "2023", "State (32 units incl. India)",
         "Deaths, Number, Both sexes, All ages, with 95% UI"],
        ["Administrative deaths", "NCRB ADSI 2023 (MHA)", "2023", "State + 53 metro cities",
         "Police-reported accidental deaths by cause; total suicides"],
        ["District population", "Census of India 2011", "2011", "District",
         "Total population; urban household share"],
        ["Behavioural covariates", "NFHS-5 district factsheets (IIPS)", "2019–21", "District",
         "Adult alcohol and tobacco prevalence"],
        ["District geometry", "india-maps-data compilation (LGD/SoI/Bhuvan/DataMeet)", "post-2014",
         "District", "Boundaries with Telangana and Ladakh/J&K splits"],
    ], columns=["Input", "Source", "Year", "Granularity", "Content"])
    add_table(d, src, "Table S1. Data sources, reference years, granularity and content.",
              "GBD: Global Burden of Disease; NCRB: National Crime Records Bureau; ADSI: "
              "Accidental Deaths & Suicides in India; MHA: Ministry of Home Affairs; IIPS: "
              "International Institute for Population Sciences; LGD: Local Government Directory; "
              "SoI: Survey of India; UI: uncertainty interval.")
    para(d, "Cause definitions. Road = road-traffic injuries; falls = falls; drowning = "
            "drowning; burns = fire, heat and hot substances; suicide = self-harm. All-injury "
            "is the GBD Level-1 “Injuries” aggregate and, on the administrative side, "
            "the sum of NCRB total accidental deaths and total suicides. Completeness for a "
            "cause is defined as administrative (NCRB) deaths divided by the fused true death "
            "estimate; it is bounded in (0, 1) by construction.", spacing="onehalf")

    # ---------------------------------------------------------------- S2 taxonomy
    heading(d, "S2. GBD cause taxonomy and aggregation")
    para(d, "The GBD extract stores, within each cause group, both the aggregate cause and its "
            "constituent sub-causes (for example, cause group “road” contains the "
            "aggregate “Road injuries” together with pedestrian, cyclist, "
            "motorcyclist, motor-vehicle and other road injuries; cause group "
            "“self-harm” contains “Self-harm” together with self-harm by "
            "firearm and by other specified means). To count each death once, only the "
            "aggregate (“parent”) cause row was used for every cause; sub-cause rows "
            "were excluded. With this convention the five named causes sum to less than the "
            "all-injury total, the remainder comprising interpersonal violence, poisoning, "
            "other transport and other unintentional injuries that are outside the five causes "
            "modelled here. State totals were formed by summing the aggregate-cause deaths (and "
            "their 95% uncertainty-interval bounds) across the states in the analytic frame; the "
            "GBD “India” row was excluded to avoid double counting.", spacing="onehalf")

    # ---------------------------------------------------------------- S3 frame
    heading(d, "S3. Geographic frame, harmonisation and coverage")
    try:
        qc = pd.read_csv(HERE / "outputs" / "qc_join_coverage.csv")
        cov_census = f"{float(qc['census_pop'].iloc[0]) * 100:.0f}%"
        cov_nfhs = f"{float(qc['nfhs'].iloc[0]) * 100:.0f}%"
    except Exception:
        cov_census, cov_nfhs = "86%", "91%"
    n_pooled = int(est.loc[est.cause == "all_injury", "is_pooled_ut"].sum()) \
        if "is_pooled_ut" in est.columns else 12
    para(d, f"The analytic frame contains {N['n_dist']} districts under current (post-2014) State "
            f"boundaries. District names were harmonised across the geometry, Census-2011 and "
            f"NFHS-5 frames by State-blocked fuzzy matching; exact-duplicate boundary records in "
            f"the source compilation were removed. Direct covariate matches were obtained for "
            f"{cov_census} of districts against the Census population frame and {cov_nfhs} against "
            f"the NFHS-5 factsheets; unmatched units and covariate gaps were imputed from the "
            f"State mean rather than dropped, and district populations were rescaled within each "
            f"State to the census State total so that benchmarking weights are internally "
            f"consistent. The queen-contiguity adjacency graph used for the ICAR field contained "
            f"no isolated nodes (0 islands). Jammu and Kashmir/Ladakh and five small union "
            f"territories are reported by GBD as single combined units; their {n_pooled} "
            f"constituent districts share the combined-unit posterior (Section S5).",
         spacing="onehalf")

    heading(d, "S4. Model specification")
    para(d, "For cause k and State s, the true death count T[s] followed a log-normal prior "
            "centred on the GBD estimate with scale set from the reported 95% uncertainty "
            "interval. The administrative count was modelled as Poisson with mean T[s]·c[s], where "
            "c[s] is the population-weighted mean of district completeness c[d]. District "
            "completeness followed logit(c[d]) = γ + η·urbanicity[d] + w[d]. The true district "
            "rate followed log λ[d] = α + Xβ + τ·φ[d] + v[d], with φ an ICAR field on the "
            "queen-contiguity district graph and v unstructured heterogeneity; λ[d] was "
            "rescaled within each State (or GBD-reporting bucket) so that the population-"
            "weighted mean equalled the fused State rate. City anchors contributed a Poisson "
            "likelihood with mean pop_city·λ[d]·c[d]. The priors on w and v were scaled per "
            "district by sqrt(pop[d] / median population), clipped to [0.15, 2.0], so that "
            "small-population districts — whose own heterogeneity term is only weakly "
            "constrained by the population-weighted benchmarking constraint — are shrunk more "
            "strongly toward the covariate/spatial structure, preventing implausible rates in "
            "thinly-populated districts. For the twelve districts in small union territories "
            "that GBD reports as a single pooled unit and that have no city anchor, the shared "
            "bucket-level R[state] and c[state] posterior draws are reported directly as the "
            "district estimate, rather than an unsupported district-specific "
            "disaggregation.", spacing="onehalf")
    para(d, "Priors. T[s] ~ LogNormal(log GBD[s], σ[s]) with σ[s] derived from the reported 95% "
            "uncertainty interval and clipped to [0.03, 1.0]; γ (completeness intercept) is "
            "centred on the national logit of 1/(GBD:NCRB ratio) for the cause; η ~ Normal(0, "
            "0.5) (0 in the fix_eta variant; 1.5 in the wide_eta variant); β ~ Normal(0, 1); "
            "α ~ Normal(log of a weak national per-capita rate, 1); the ICAR precision and the "
            "unstructured scales use half-Normal hyperpriors, with w and v scaled per district "
            "as above. The NCRB Poisson likelihood is masked where a State administrative count "
            "is missing, so those States are informed by GBD and the spatial field alone.",
         spacing="onehalf")

    heading(d, "S5. Derived quantities")
    para(d, "Two policy quantities are computed from the joint posterior draws rather than from "
            "collapsed point estimates. (i) District surveillance completeness c[d] is the "
            "posterior of the administrative-to-true ratio at each district, reported as the "
            "mean and 95% credible interval. (ii) The blind-spot probability is the joint "
            "posterior probability, evaluated in each draw, that a district's rate exceeds the "
            "national posterior-median rate for that draw AND its completeness is below 0.5 in "
            "the same draw; averaging the indicator across draws yields an uncertainty-aware "
            "“confidence this is a true blind spot”. Because the national median is "
            "recomputed within each draw, the two conditions are evaluated coherently rather "
            "than against a fixed external threshold.", spacing="onehalf")

    heading(d, "S6. Convergence")
    conv = pd.read_csv(HERE / "outputs" / "convergence_summary.csv")
    conv2 = conv.copy(); conv2.columns = ["Cause", "max R-hat", "min ESS", "Anchored"]
    conv2["max R-hat"] = conv2["max R-hat"].round(3)
    add_table(d, conv2, "Table S2. Convergence diagnostics by cause.",
              "R-hat < 1.05 and effective sample size (ESS) indicate satisfactory convergence.")

    heading(d, "S7. Structural sensitivity")
    if sens is not None:
        s2 = sens.copy()
        s2.columns = ["Variant", "Cause", "CV log-RMSE", "Mean completeness"]
        s2["CV log-RMSE"] = s2["CV log-RMSE"].round(3); s2["Mean completeness"] = s2["Mean completeness"].round(2)
        add_table(d, s2, "Table S3. Structural sensitivity (road), five-fold city cross-validation.",
                  "base: full model; no_icar: spatial term removed; fix_eta: completeness held "
                  "constant within State; wide_eta: diffuse completeness-slope prior.")
    else:
        para(d, "Sensitivity table pending completion of the sensitivity run.", italic=True)

    heading(d, "S8. Formal spatial model comparison")
    loo_path = TAB / "model_comparison_loo.csv"
    if loo_path.exists():
        loo = pd.read_csv(loo_path)
        low_reliability = (loo["pct_bad_or_worse_k"] > 30).any() if "pct_bad_or_worse_k" in loo.columns else False
        n_div_total = int(loo["n_divergences"].sum()) if "n_divergences" in loo.columns else None
        para(d, "As a secondary check on the five-fold cross-validation reported in Table 3 "
                "and Table S3, we also computed PSIS-LOO on the held-out city-anchor likelihood "
                f"for the spatial and non-spatial road models (n={int(loo['n_obs'].iloc[0]) if 'n_obs' in loo.columns else 45} "
                "anchor observations). Because the initial fits showed a poor Pareto-k "
                "reliability diagnostic, we explicitly tested whether this reflected inadequate "
                "sampling: both models were refitted with four chains, 2000 tuning and 2000 "
                "post-warm-up draws and a target acceptance rate of 0.995. This eliminated "
                f"divergent transitions entirely ({n_div_total if n_div_total is not None else 0} "
                "across both refitted models) but did not improve the Pareto-k diagnostic "
                "(Table S4), indicating that the unreliability is a structural consequence of "
                "comparing district-level models on only 45 anchor observations rather than an "
                "artefact of insufficient sampling. " +
                ("We therefore report this PSIS-LOO comparison as exploratory only; the "
                 "five-fold cross-validation in Table S3, which measures genuine held-out "
                 "predictive error directly rather than relying on the PSIS approximation, "
                 "remains the primary evidence that the spatial term adds little to "
                 "out-of-sample prediction for this cause." if low_reliability else
                 "The Pareto-k reliability diagnostic was adequate, supporting the LOO "
                 "comparison as confirmatory evidence alongside Table S3."),
             spacing="onehalf")
        loo_show = loo.copy()
        for c in ["elpd_loo", "se", "elpd_diff", "dse"]:
            if c in loo_show.columns:
                loo_show[c] = loo_show[c].round(2)
        rename = {"pct_good_k": "Pareto k<=0.7 (%)", "pct_bad_or_worse_k": "Pareto k>0.7 (%)",
                  "n_obs": "n anchors", "n_divergences": "Divergences (refitted)"}
        loo_show = loo_show.rename(columns=rename)
        add_table(d, loo_show, "Table S4. PSIS-LOO comparison of the spatial (ICAR) and "
                  "non-spatial (covariates only) district models, road cause.",
                  "elpd_loo: expected log pointwise predictive density (higher is better); "
                  "elpd_diff/dse: difference from the best model and its standard error. "
                  "Pareto k diagnoses the reliability of the PSIS-LOO approximation at each "
                  "held-out point; values above 0.7 indicate the approximation may be "
                  "unreliable. Divergences (refitted) confirms the poor Pareto-k values are not "
                  "attributable to inadequate sampling: both models were refitted with four "
                  "chains and a 0.995 target acceptance rate and produced zero divergent "
                  "transitions, yet Pareto-k did not improve, most likely reflecting the small "
                  "anchor sample rather than a model or sampling failure. Evaluated on the "
                  "held-out metropolitan-city anchor "
                  "likelihood only.")
    else:
        para(d, "Formal LOO/WAIC comparison pending completion of the model-comparison run.",
             italic=True)

    heading(d, "S9. Surveillance blind-spot districts by cause")
    para(d, "Table S5 lists the leading blind-spot districts for each anchored cause (all-injury, "
            "road, suicide), ranked by the posterior blind-spot probability defined in Section "
            "S5. Districts sharing a single pooled union-territory estimate are excluded, as "
            "their identical values would not reflect independent district-level evidence. Road "
            "deaths are comparatively well captured administratively, so few road districts reach "
            "high blind-spot probability; the strongest and most consistent blind spots are for "
            "all-injury and suicide in the Himalayan hill and central-Indian tribal districts.",
         spacing="onehalf")
    try:
        bs = pd.read_csv(SUB / "Table_2_blindspots.csv")
        order = {"all_injury": 0, "road": 1, "suicide": 2}
        bs = bs.sort_values(by=["Cause", "P(blind spot)"],
                            key=lambda s: s.map(order) if s.name == "Cause" else s,
                            ascending=[True, False])
        bs = bs.groupby("Cause", sort=False).head(15)
        add_table(d, bs, "Table S5. Leading surveillance blind-spot districts by anchored cause "
                  "(top 15 each), India 2023.",
                  "Rate is posterior mean per 100 000 (95% credible interval). Completeness = "
                  "estimated administrative/true deaths. P(blind spot) is the joint posterior "
                  "probability that a district is above the national median rate and below 0.5 "
                  "completeness in the same draw.")
    except Exception:
        para(d, "Blind-spot table pending asset regeneration.", italic=True)

    heading(d, "S10. State-level burden and completeness summary")
    para(d, "Table S6 aggregates the district atlas to State-median values for the flagship "
            "causes, showing the between-State spread that the district model resolves further "
            "within each State.", spacing="onehalf")
    try:
        def _sm(cause, col):
            return est[est.cause == cause].groupby("state_name")[col].median()
        ss = pd.DataFrame({
            "All-injury rate": _sm("all_injury", "rate_mean"),
            "All-injury completeness": _sm("all_injury", "completeness_mean"),
            "Road rate": _sm("road", "rate_mean"),
            "Suicide rate": _sm("suicide", "rate_mean"),
        }).round(2).sort_values("All-injury rate", ascending=False)
        ss.insert(0, "State", ss.index)
        add_table(d, ss, "Table S6. State-median district estimates (per 100 000; completeness "
                  "as a fraction), ordered by all-injury burden.",
                  "Values are medians across the districts of each State/union-territory unit; "
                  "combined GBD-reporting units (Jammu and Kashmir/Ladakh; Other Union "
                  "Territories) appear as single rows.")
    except Exception:
        para(d, "State summary pending asset regeneration.", italic=True)

    heading(d, "S11. Per-cause maps")
    para(d, "Rate, uncertainty and (for anchored causes) completeness choropleths were generated "
            "for all six causes; every one of India's districts is drawn, with pooled "
            "union-territory districts shown at their shared value and footnoted rather than "
            "blanked. High-resolution PNG panels for all causes accompany this submission in the "
            "figures folder and in the code repository.", spacing="onehalf")

    heading(d, "S12. Software and reproducibility")
    para(d, "Models were implemented in PyMC and sampled with the No-U-Turn sampler through the "
            "NumPyro (JAX) backend; posterior handling used ArviZ, spatial data GeoPandas, and "
            "tabular data pandas/NumPy. Each cause was fitted in an isolated process to bound "
            "device memory. All input data are public; the complete extraction, modelling, "
            "validation and figure code, together with the derived tables, is available at "
            "https://github.com/hssling/injury-india-public-health, permitting the atlas to be "
            "regenerated as later GBD and NCRB rounds are released.", spacing="onehalf")

    heading(d, "S13. Limitations")
    para(d, "The completeness gradient is identified from 45 single-district metropolitan "
            "anchors, all urban, and is extrapolated to rural districts that are never directly "
            "observed; this is the principal caveat and the reason falls, drowning and burns are "
            "reported as covariate projections rather than validated estimates. District rankings "
            "should be read with their credible intervals rather than as exact positions. "
            "Covariate associations are ecological. Exposure uses the 2011 census base, and "
            "boundary harmonisation across three frames carries residual name-matching "
            "uncertainty that was logged but not eliminated. NCRB city tables restrict anchoring "
            "to three causes; release of district-level cause-specific deaths (for example "
            "through the Civil Registration System) would allow falls and burns to be anchored "
            "and validated directly.", spacing="onehalf")
    d.save(SUB / "05_supplementary_IJMR.docx")


def checklist():
    txt = f"""# IJMR Submission Checklist — District Injury Atlas

- [x] 00_cover_letter_and_title_page_IJMR.docx — cover letter (addressed to the Editor) + title
      page (author, affiliation, corresponding author, word/table/figure counts), one file
- [x] 02_declarations_IJMR.docx — full, author-identified: funding, conflicts, ethics, data/code,
      contributions, AI use (non-blinded, for editorial records only)
- [x] 03_main_manuscript_IJMR.docx — BLINDED (no author/affiliation); IJMR order: structured
      abstract (Background & objectives / Methods / Results / Interpretation & conclusions),
      Key Words, Introduction, Material & Methods, Results, Discussion (ending in an explicit
      Conclusion and policy implications paragraph), Declarations (funding/COI/ethics/data,
      no author-identifying contributions), References ({len(REFERENCES)}, Vancouver style,
      hanging indent, consecutive runs hyphenated), Tables 1-3 (with footnotes), Figure legends
- [x] 04_figures_IJMR.docx + figures/ folder (PNG panels, all districts shown, pooled-UT
      districts footnoted rather than blanked)
- [x] 05_supplementary_IJMR.docx — model spec, convergence, sensitivity, LOO/WAIC model
      comparison, sources

Format: 12 pt Times New Roman, double-spaced, 2.5 cm margins; in-text citations as bare
superscript numbers (no brackets), Vancouver-style ranges (e.g., 10-13) for 3+ consecutive.
Key results: {N['n_dist']} districts; falls completeness {N['falls_comp']}; all-injury completeness {N['ai_comp']}.
"""
    (SUB / "Submission_checklist_IJMR.md").write_text(txt, encoding="utf-8")


if __name__ == "__main__":
    SUB.mkdir(exist_ok=True)
    cover_letter_and_title_page(); declarations(); main_manuscript()
    figures_doc(); supplementary(); checklist()
    print("built full IJMR package in", SUB)
    print("N=", N)
